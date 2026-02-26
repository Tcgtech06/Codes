import customtkinter as ct
import os
import sys
import json
import shutil
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import filedialog, messagebox
import base64
from PIL import Image, ImageTk
import re

# --- CONFIGURATION ---
# Get the correct base path for both development and frozen (EXE) environments
if getattr(sys, 'frozen', False):
    # Running as compiled EXE
    BASE_DIR = os.path.dirname(sys.executable)
    # Use AppData for writable files
    APPDATA_DIR = os.getenv('APPDATA')
    DB_FOLDER = os.path.join(APPDATA_DIR, "TCG_Invoice", "DB")
    # Resources (fonts, images) from installation directory
    RESOURCE_DIR = os.path.join(BASE_DIR, "_internal") if os.path.exists(os.path.join(BASE_DIR, "_internal")) else BASE_DIR
else:
    # Running as Python script
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    DB_FOLDER = os.path.join(BASE_DIR, "DB")
    RESOURCE_DIR = BASE_DIR

# File paths
LOGOS_FOLDER = os.path.join(DB_FOLDER, "logos")
HISTORY_FILE = os.path.join(DB_FOLDER, "history.json")
PROFILES_FILE = os.path.join(DB_FOLDER, "profiles.json")
LICENSE_FILE = os.path.join(DB_FOLDER, "license.json")  # New license file

# Ensure directories exist
os.makedirs(DB_FOLDER, exist_ok=True)
os.makedirs(LOGOS_FOLDER, exist_ok=True)

# License keys - dynamic with current year
current_year = str(datetime.now().year)
VALID_LICENSE_KEYS = {
    "demo": f"tcgtech{current_year}d",
    "subscription": f"tcgtech{current_year}s", 
    "permanent": f"tcgtech{current_year}p"
}

# Default Profiles Structure
DEFAULT_PROFILES = {
    "biz_1": {
        "name": "Tech Solutions Inc.",
        "address": "123 Silicon Valley, CA",
        "email": "contact@techsolutions.com",
        "phone": "+1 (555) 123-4567",
        "color": "#1e40af",  # Blue
        "icon": "",
        "style": "Modern",
        "logo": None,
        "last_invoice_num": 1000
    },
    "biz_2": {
        "name": "Green Earth Gardens",
        "address": "45 Nature Way, Oregon",
        "email": "info@greenearth.com",
        "phone": "+1 (555) 987-6543",
        "color": "#059669",  # Green
        "icon": "",
        "style": "Classic",
        "logo": None,
        "last_invoice_num": 2000
    },
    "biz_3": {
        "name": "Express Logistics",
        "address": "88 Cargo Blvd, NY",
        "email": "ops@expresslogistics.com",
        "phone": "+1 (555) 456-7890",
        "color": "#ec4899",  # Pink
        "icon": "",
        "style": "Industrial",
        "logo": None,
        "last_invoice_num": 3000
    },
    "biz_4": {
        "name": "Creative Studio",
        "address": "99 Design Ave, London",
        "email": "hello@creativestudio.uk",
        "phone": "+44 20 7123 4567",
        "color": "#7c3aed",  # Purple
        "icon": "",
        "style": "Minimalist",
        "logo": None,
        "last_invoice_num": 4000
    }
}


class SplashScreen(ct.CTkToplevel):
    def __init__(self, parent):
        super().__init__(parent)
        
        # Configure splash window - FULL SCREEN
        self.title("")
        self.resizable(False, False)
        
        # Remove window decorations
        self.overrideredirect(True)
        
        # Get screen dimensions and make fullscreen
        self.update_idletasks()
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        self.geometry(f'{screen_width}x{screen_height}+0+0')
        
        # Configure background - WHITE
        self.configure(fg_color="white")
        
        # Main container - centered
        container = ct.CTkFrame(self, fg_color="white")
        container.pack(expand=True, fill="both")
        
        # Application Logo with fade-in animation - BIGGER SIZE
        self.logo_label = ct.CTkLabel(container, text="", fg_color="white")
        self.logo_label.pack(pady=(40, 20))
        
        try:
            logo_img = Image.open(os.path.join(RESOURCE_DIR, "Image", "logo.png"))
            logo_img = logo_img.resize((250, 250), Image.Resampling.LANCZOS)
            logo_photo = ct.CTkImage(light_image=logo_img, dark_image=logo_img, size=(250, 250))
            self.logo_label.configure(image=logo_photo)
            self.logo_label.image = logo_photo
        except Exception as e:
            print(f"Could not load logo: {e}")
            self.logo_label.configure(text="??", font=ct.CTkFont(size=120))
        
        # "A PRODUCT OF" text in BOLD below logo
        ct.CTkLabel(container, text="A PRODUCT OF", 
                   font=ct.CTkFont(size=24, weight="bold"),
                   text_color="#1e40af", fg_color="white").pack(pady=(20, 30))
        
        # TCG Tech Logo - original size, no stretching
        self.tcg_label = ct.CTkLabel(container, text="", fg_color="white")
        self.tcg_label.pack(pady=20)
        
        tcg_loaded = False
        try:
            # Check for TCG Tech logo
            tcg_path = os.path.join(RESOURCE_DIR, "Image", "Tcgtech.png")
            
            if tcg_path and os.path.exists(tcg_path):
                print(f"Loading TCG logo from: {tcg_path}")
                tcg_img = Image.open(tcg_path)
                # Increase size - scale to 300x120 max
                tcg_img.thumbnail((300, 120), Image.Resampling.LANCZOS)
                tcg_photo = ct.CTkImage(light_image=tcg_img, dark_image=tcg_img, size=tcg_img.size)
                self.tcg_label.configure(image=tcg_photo)
                self.tcg_label.image = tcg_photo
                tcg_loaded = True
                print("TCG logo loaded successfully")
            else:
                print(f"TCG logo file not found at: {tcg_path}")
        except Exception as e:
            print(f"Error loading TCG Tech logo: {e}")
            import traceback
            traceback.print_exc()
        
        # Fallback if image didn't load
        if not tcg_loaded:
            ct.CTkLabel(container, text="TCG TECH", 
                       font=ct.CTkFont(size=28, weight="bold"),
                       text_color="#1e40af", fg_color="white").pack(pady=20)
        
        # Start fade-in animation
        self.alpha = 0.0
        self.animate_fade_in()
        
        # Close splash screen after 4 seconds
        self.after(4000, self.close_splash)
    
    def animate_fade_in(self):
        """Animate fade-in effect"""
        if self.alpha < 1.0:
            self.alpha += 0.05
            self.attributes('-alpha', self.alpha)
            self.after(30, self.animate_fade_in)
    
    def close_splash(self):
        self.destroy()


class LoginScreen(ct.CTkToplevel):
    def __init__(self, parent):
        super().__init__(parent)

        # Configure login window
        self.title("TCG Tech Invoice Generator - Login")
        self.resizable(False, False)
        self.geometry("500x600")
        self.configure(fg_color="white")

        # Center the window
        self.update_idletasks()
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        x = (screen_width - 500) // 2
        y = (screen_height - 600) // 2
        self.geometry(f"500x600+{x}+{y}")

        # Remove window decorations
        self.overrideredirect(True)

        # Main container
        container = ct.CTkFrame(self, fg_color="white", corner_radius=20)
        container.pack(expand=True, fill="both", padx=30, pady=30)

        # TCG Logo at top
        self.tcg_label = ct.CTkLabel(container, text="", fg_color="white")
        self.tcg_label.pack(pady=(20, 10))

        try:
            # Check for TCG Tech logo
            tcg_path = os.path.join(RESOURCE_DIR, "Image", "Tcgtech.png")

            if tcg_path and os.path.exists(tcg_path):
                print(f"Loading TCG logo from: {tcg_path}")
                tcg_img = Image.open(tcg_path)
                # Increase size - scale to 200x80 max
                tcg_img.thumbnail((200, 80), Image.Resampling.LANCZOS)
                tcg_photo = ct.CTkImage(light_image=tcg_img, dark_image=tcg_img, size=tcg_img.size)
                self.tcg_label.configure(image=tcg_photo)
                self.tcg_label.image = tcg_photo
                print("TCG logo loaded successfully")
            else:
                print(f"TCG logo file not found at: {tcg_path}")
        except Exception as e:
            print(f"Error loading TCG Tech logo: {e}")
            import traceback
            traceback.print_exc()

        # Title
        ct.CTkLabel(container, text="INVOICE GENERATOR",
                   font=ct.CTkFont(size=24, weight="bold"),
                   text_color="#1e40af", fg_color="white").pack(pady=(10, 5))
        ct.CTkLabel(container, text="Choose Your Account Type",
                   font=ct.CTkFont(size=16, weight="bold"),
                   text_color="#64748b", fg_color="white").pack(pady=(0, 30))

        # Account Type Selection Frame
        selection_frame = ct.CTkFrame(container, fg_color="#f8fafc", corner_radius=15)
        selection_frame.pack(fill="x", padx=20, pady=(0, 20))

        # Demo Account Button
        self.demo_btn = ct.CTkButton(
            selection_frame,
            text="?? Demo Account",
            command=self.open_demo_window,
            height=60,
            font=ct.CTkFont(size=16, weight="bold"),
            fg_color="#10b981",
            hover_color="#059669",
            corner_radius=12
        )
        self.demo_btn.pack(fill="x", padx=20, pady=(20, 10))

        # Professional Account Button
        self.pro_btn = ct.CTkButton(
            selection_frame,
            text="?? Professional Account",
            command=self.open_professional_window,
            height=60,
            font=ct.CTkFont(size=16, weight="bold"),
            fg_color="#3b82f6",
            hover_color="#2563eb",
            corner_radius=12
        )
        self.pro_btn.pack(fill="x", padx=20, pady=(10, 20))

        # Dynamic content area (initially hidden)
        self.content_frame = ct.CTkFrame(container, fg_color="transparent")
        self.content_frame.pack(fill="x", padx=20, pady=20)
        self.content_frame.pack_forget()  # Hide initially

        # Footer
        footer_frame = ct.CTkFrame(container, fg_color="transparent")
        footer_frame.pack(fill="x", pady=(20, 0))

        ct.CTkLabel(footer_frame, text=f"� {datetime.now().year} TCG Tech. All rights reserved.",
                   font=ct.CTkFont(size=10),
                   text_color="#94a3b8", fg_color="transparent").pack()

        # State variables
        self.selected_account_type = None

    def open_demo_window(self):
        """Open a separate window for demo account license key entry"""
        # Hide current window temporarily
        self.withdraw()

        # Create demo window
        demo_window = ct.CTkToplevel(self)
        demo_window.title("TCG Tech - Demo Account")
        demo_window.geometry("450x400")
        demo_window.resizable(False, False)
        demo_window.configure(fg_color="white")

        # Center the window
        demo_window.update_idletasks()
        screen_width = demo_window.winfo_screenwidth()
        screen_height = demo_window.winfo_screenheight()
        x = (screen_width - 450) // 2
        y = (screen_height - 400) // 2
        demo_window.geometry(f"450x400+{x}+{y}")

        demo_window.overrideredirect(True)
        demo_window.transient(self)
        demo_window.grab_set()

        # Main container
        container = ct.CTkFrame(demo_window, fg_color="white", corner_radius=20)
        container.pack(expand=True, fill="both", padx=30, pady=30)

        # Title
        ct.CTkLabel(container, text="Demo Account",
                   font=ct.CTkFont(size=24, weight="bold"),
                   text_color="#059669", fg_color="white").pack(pady=(20, 10))

        # Subtitle
        ct.CTkLabel(container, text="Enter your demo license key",
                   font=ct.CTkFont(size=14),
                   text_color="#64748b", fg_color="white").pack(pady=(0, 30))

        # License Key Input
        input_frame = ct.CTkFrame(container, fg_color="#f8fafc", corner_radius=10)
        input_frame.pack(fill="x", pady=(0, 20))

        ct.CTkLabel(input_frame, text="Demo License Key:",
                   font=ct.CTkFont(size=14, weight="bold"),
                   text_color="#374151").pack(pady=(20, 10))

        license_key_entry = ct.CTkEntry(
            input_frame,
            placeholder_text="Enter your demo license key...",
            height=45,
            corner_radius=8,
            border_width=2,
            font=ct.CTkFont(size=12)
        )
        license_key_entry.pack(fill="x", padx=20, pady=(0, 20))

        # Buttons
        button_frame = ct.CTkFrame(container, fg_color="transparent")
        button_frame.pack(fill="x", pady=(10, 0))

        def activate_account():
            license_key = license_key_entry.get().strip()
            if not license_key:
                messagebox.showerror("Error", "Please enter a demo license key.")
                return

            # Validate against expected keys
            expected_key = VALID_LICENSE_KEYS.get("demo")
            if license_key != expected_key:
                messagebox.showerror("Error", "Invalid license key for demo account.")
                return

            # Create license data and activate
            license_data = {
                "account_type": "demo",
                "license_key": license_key,
                "activation_date": datetime.now().isoformat(),
                "is_active": True
            }

            # Save license data using InvoiceApp method
            temp_app = InvoiceApp.__new__(InvoiceApp)  # Create temp instance
            temp_app.save_license_data(license_data)

            messagebox.showinfo("Demo Activated",
                              "Demo activated successfully!\n\nLicense Key: {}\n\nWelcome to TCG Tech Invoice Generator Demo.".format(license_key))

            demo_window.destroy()
            self.proceed_to_main_app()

        def back_to_login():
            demo_window.destroy()
            self.deiconify()  # Show main login window again

        ct.CTkButton(button_frame, text="?? Activate Demo",
                    command=activate_account,
                    height=45,
                    fg_color="#10b981",
                    hover_color="#059669",
                    font=ct.CTkFont(size=14, weight="bold")).pack(side="left", expand=True, padx=(0, 5))

        ct.CTkButton(button_frame, text="?? Back",
                    command=back_to_login,
                    height=45,
                    fg_color="#64748b",
                    hover_color="#475569",
                    font=ct.CTkFont(size=14, weight="bold")).pack(side="right", expand=True, padx=(5, 0))

    def open_professional_window(self):
        """Open a separate window for professional account options"""
        # Hide current window temporarily
        self.withdraw()

        # Create professional window
        pro_window = ct.CTkToplevel(self)
        pro_window.title("TCG Tech - Professional Account")
        pro_window.geometry("450x400")
        pro_window.resizable(False, False)
        pro_window.configure(fg_color="white")

        # Center the window
        pro_window.update_idletasks()
        screen_width = pro_window.winfo_screenwidth()
        screen_height = pro_window.winfo_screenheight()
        x = (screen_width - 450) // 2
        y = (screen_height - 400) // 2
        pro_window.geometry(f"450x400+{x}+{y}")

        pro_window.overrideredirect(True)
        pro_window.transient(self)
        pro_window.grab_set()

        # Main container
        container = ct.CTkFrame(pro_window, fg_color="white", corner_radius=20)
        container.pack(expand=True, fill="both", padx=30, pady=30)

        # Title
        ct.CTkLabel(container, text="Professional Account",
                   font=ct.CTkFont(size=24, weight="bold"),
                   text_color="#2563eb", fg_color="white").pack(pady=(20, 10))

        # Subtitle
        ct.CTkLabel(container, text="Choose your account type",
                   font=ct.CTkFont(size=14),
                   text_color="#64748b", fg_color="white").pack(pady=(0, 30))

        # Options
        options_frame = ct.CTkFrame(container, fg_color="#f8fafc", corner_radius=10)
        options_frame.pack(fill="x", pady=(0, 20))

        def select_subscription():
            pro_window.destroy()
            self.open_license_activation("subscription")

        def select_permanent():
            pro_window.destroy()
            self.open_license_activation("permanent")

        def back_to_login():
            pro_window.destroy()
            self.deiconify()  # Show main login window again

        # Subscription Account Button
        ct.CTkButton(
            options_frame,
            text="?? Subscription Account",
            command=select_subscription,
            height=50,
            font=ct.CTkFont(size=14, weight="bold"),
            fg_color="#8b5cf6",
            hover_color="#7c3aed",
            corner_radius=10
        ).pack(fill="x", padx=20, pady=(20, 10))

        # Permanent Account Button
        ct.CTkButton(
            options_frame,
            text="?? Permanent Account",
            command=select_permanent,
            height=50,
            font=ct.CTkFont(size=14, weight="bold"),
            fg_color="#f59e0b",
            hover_color="#d97706",
            corner_radius=10
        ).pack(fill="x", padx=20, pady=(10, 20))

        # Back button
        ct.CTkButton(container, text=" Back to Login",
                    command=back_to_login,
                    height=45,
                    fg_color="#64748b",
                    hover_color="#475569",
                    font=ct.CTkFont(size=14, weight="bold")).pack(fill="x", pady=(10, 0))

    def open_license_activation(self, account_type):
        """Open license activation window for subscription and permanent accounts"""
        # Create activation window
        activation_window = ct.CTkToplevel(self)
        activation_window.title(f"TCG Tech - {account_type.title()} Account Activation")
        activation_window.geometry("450x400")
        activation_window.resizable(False, False)
        activation_window.configure(fg_color="white")

        # Center the window
        activation_window.update_idletasks()
        screen_width = activation_window.winfo_screenwidth()
        screen_height = activation_window.winfo_screenheight()
        x = (screen_width - 450) // 2
        y = (screen_height - 400) // 2
        activation_window.geometry(f"450x400+{x}+{y}")

        activation_window.overrideredirect(True)
        activation_window.transient(self)
        activation_window.grab_set()

        # Main container
        container = ct.CTkFrame(activation_window, fg_color="white", corner_radius=20)
        container.pack(expand=True, fill="both", padx=30, pady=30)

        # Title based on account type
        if account_type == "subscription":
            title_text = "Subscription Account"
            subtitle_text = "Enter your subscription license key"
            button_text = " Activate Subscription"
            color = "#7c3aed"
        elif account_type == "permanent":
            title_text = "Permanent Account"
            subtitle_text = "Enter your permanent license key"
            button_text = " Activate Permanent"
            color = "#d97706"

        ct.CTkLabel(container, text=title_text,
                   font=ct.CTkFont(size=24, weight="bold"),
                   text_color=color, fg_color="white").pack(pady=(20, 10))

        ct.CTkLabel(container, text=subtitle_text,
                   font=ct.CTkFont(size=14),
                   text_color="#64748b", fg_color="white").pack(pady=(0, 30))

        # License Key Input
        input_frame = ct.CTkFrame(container, fg_color="#f8fafc", corner_radius=10)
        input_frame.pack(fill="x", pady=(0, 20))

        ct.CTkLabel(input_frame, text=f"{title_text} License Key:",
                   font=ct.CTkFont(size=14, weight="bold"),
                   text_color="#374151").pack(pady=(20, 10))

        license_key_entry = ct.CTkEntry(
            input_frame,
            placeholder_text="Enter your license key...",
            height=45,
            corner_radius=8,
            border_width=2,
            font=ct.CTkFont(size=12)
        )
        license_key_entry.pack(fill="x", padx=20, pady=(0, 20))

        # Buttons
        button_frame = ct.CTkFrame(container, fg_color="transparent")
        button_frame.pack(fill="x", pady=(10, 0))

        def activate_account():
            license_key = license_key_entry.get().strip()
            if not license_key:
                messagebox.showerror("Error", f"Please enter a {account_type} license key.")
                return

            # Validate against expected keys
            expected_key = VALID_LICENSE_KEYS.get(account_type)
            if license_key != expected_key:
                messagebox.showerror("Error", f"Invalid license key for {account_type} account.")
                return

            # Create license data and activate
            license_data = {
                "account_type": account_type,
                "license_key": license_key,
                "activation_date": datetime.now().isoformat(),
                "is_active": True
            }

            # Save license data using InvoiceApp method
            temp_app = InvoiceApp.__new__(InvoiceApp)  # Create temp instance
            temp_app.save_license_data(license_data)

            messagebox.showinfo(f"{title_text} Activated",
                              f"{title_text} activated successfully!\n\nLicense Key: {license_key}\n\nWelcome to TCG Tech Invoice Generator {title_text}!")

            activation_window.destroy()
            self.deiconify()  # Show main login window again

        def back_to_previous():
            activation_window.destroy()
            # Show the professional window again
            self.open_professional_window()

        ct.CTkButton(button_frame, text=button_text,
                    command=activate_account,
                    height=45,
                    fg_color=color,
                    hover_color=color.replace("#", "#").replace("0", "3").replace("5", "6").replace("6", "9").replace("9", "b"),  # Darker shade
                    font=ct.CTkFont(size=14, weight="bold")).pack(side="left", expand=True, padx=(0, 5))

        ct.CTkButton(button_frame, text=" Back",
                    command=back_to_previous,
                    height=45,
                    fg_color="#64748b",
                    hover_color="#475569",
                    font=ct.CTkFont(size=14, weight="bold")).pack(side="right", expand=True, padx=(5, 0))

    def proceed_to_main_app(self):
        """Close login screen and proceed to main application"""
        self.destroy()


class InvoiceApp(ct.CTk):
    def __init__(self):
        super().__init__()

        self.title("Tcg Tech Multi Business Invoice")
        self.geometry("1400x950")
        ct.set_appearance_mode("light")
        ct.set_default_color_theme("blue")

        # Set window icon
        try:
            icon_img = Image.open(os.path.join(RESOURCE_DIR, "Image", "logo.png"))
            icon_photo = ImageTk.PhotoImage(icon_img)
            self.iconphoto(True, icon_photo)
            self._icon_ref = icon_photo  # Keep reference
        except Exception as e:
            print(f"Could not set window icon: {e}")

        # Load Data
        self.profiles = self.load_profiles()
        self.history = self.load_history()
        self.items = []

        # State: Default to Business 1
        self.current_biz_id = "biz_1"
        self.current_theme = self.profiles[self.current_biz_id]["color"]
        
        # Draggable containers system
        self.draggable_containers = {}  # Track all draggable containers
        self.drag_data = {"x": 0, "y": 0, "item": None}  # Track drag state

        # Layout
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        self.setup_app_logo()

        self.setup_sidebar()
        self.setup_main_content()

        # Apply initial theme
        self.refresh_theme()

    def create_draggable_container(self, parent, container_id, container_name):
        """Create a wrapper with drag handle for a container"""
        # Outer wrapper
        wrapper = ct.CTkFrame(parent, fg_color="transparent")
        wrapper.pack(fill="x", padx=30, pady=(25, 0))
        
        # Drag handle header
        drag_header = ct.CTkFrame(wrapper, fg_color=("#3b82f6", "#2563eb"),
                                 corner_radius=8, height=32)
        drag_header.pack(fill="x", pady=(0, 5))
        drag_header.pack_propagate(False)
        
        header_content = ct.CTkFrame(drag_header, fg_color="transparent")
        header_content.pack(fill="x", padx=12, pady=8)
        
        # Drag handle with three lines
        drag_label = ct.CTkLabel(header_content, text=container_name.upper(),
                                font=ct.CTkFont(size=11, weight="bold"),
                                text_color="white")
        drag_label.pack(side="left", expand=True, anchor="w")
        
        # Context menu on double-click
        def on_double_click(event):
            self.show_container_context_menu(container_id, container_name, wrapper, drag_label)
        
        drag_label.bind("<Double-Button-1>", on_double_click)
        drag_header.bind("<Double-Button-1>", on_double_click)
        
        # Drag functionality
        def on_drag_start(event):
            self.drag_data["x"] = event.x_root
            self.drag_data["y"] = event.y_root
            self.drag_data["item"] = (container_id, wrapper)
            drag_header.configure(fg_color=("#2563eb", "#1d4ed8"))
        
        def on_drag_motion(event):
            if self.drag_data["item"]:
                delta_y = event.y_root - self.drag_data["y"]
                current_y = wrapper.winfo_y()
                self.draggable_containers[container_id]["wrapper"].place(
                    x=30, y=current_y + delta_y, width=wrapper.winfo_width()
                )
                self.drag_data["y"] = event.y_root
        
        def on_drag_release(event):
            if self.drag_data["item"]:
                drag_header.configure(fg_color=("#3b82f6", "#2563eb"))
                messagebox.showinfo("Layout Updated", 
                                  f"Container '{container_name}' position updated!\nLayout changes saved.")
                self.drag_data["item"] = None
        
        drag_header.bind("<Button-1>", on_drag_start)
        drag_header.bind("<B1-Motion>", on_drag_motion)
        drag_header.bind("<ButtonRelease-1>", on_drag_release)
        drag_label.bind("<Button-1>", on_drag_start)
        drag_label.bind("<B1-Motion>", on_drag_motion)
        drag_label.bind("<ButtonRelease-1>", on_drag_release)
        
        # Store container data
        self.draggable_containers[container_id] = {
            "wrapper": wrapper,
            "header": drag_header,
            "name": container_name,
            "label": drag_label
        }
        
        return wrapper
    
    def show_container_context_menu(self, container_id, container_name, wrapper, drag_label):
        """Show context menu for container rename and delete"""
        # Create popup window
        popup = ct.CTkToplevel(self)
        popup.title("Container Options")
        popup.geometry("300x150")
        popup.transient(self)
        popup.grab_set()
        
        # Title
        ct.CTkLabel(popup, text=f"Edit: {container_name}",
                   font=ct.CTkFont(size=14, weight="bold"),
                   text_color=("#1e40af", "#60a5fa")).pack(pady=15)
        
        # Rename option
        def rename_container():
            rename_dialog = ct.CTkInputDialog(
                text=f"Enter new name for '{container_name}':",
                title="Rename Container"
            )
            new_name = rename_dialog.get_input()
            if new_name:
                self.draggable_containers[container_id]["name"] = new_name
                drag_label.configure(text="= " + new_name.upper())
                messagebox.showinfo("Renamed", f"Container renamed to '{new_name}'!")
                self.save_profiles()
            popup.destroy()
        
        ct.CTkButton(popup, text="?? Rename Container",
                    command=rename_container,
                    width=250, height=40,
                    fg_color=("#8b5cf6", "#7c3aed"),
                    hover_color=("#7c3aed", "#6d28d9"),
                    font=ct.CTkFont(size=12, weight="bold")).pack(pady=5, padx=15)
        
        # Delete option
        def delete_container():
            if messagebox.askyesno("Delete", f"Remove '{container_name}' from layout?\n\nThis action cannot be undone."):
                wrapper.destroy()
                if container_id in self.draggable_containers:
                    del self.draggable_containers[container_id]
                messagebox.showinfo("Deleted", f"Container '{container_name}' removed!")
                self.save_profiles()
            popup.destroy()
        
        ct.CTkButton(popup, text="??? Delete Container",
                    command=delete_container,
                    width=250, height=40,
                    fg_color=("#ef4444", "#dc2626"),
                    hover_color=("#dc2626", "#b91c1c"),
                    font=ct.CTkFont(size=12, weight="bold")).pack(pady=5, padx=15)

    def setup_app_logo(self):
        """Add application logo to the top left corner of the sidebar"""
        # This will be called before setup_sidebar, so we'll add it in setup_sidebar instead
        pass

    def load_profiles(self):
        if os.path.exists(PROFILES_FILE):
            try:
                with open(PROFILES_FILE, 'r') as f:
                    data = json.load(f)
                    # Merge with defaults to ensure structure exists if file is old
                    for key, val in DEFAULT_PROFILES.items():
                        if key not in data:
                            data[key] = val
                    return data
            except:
                return DEFAULT_PROFILES
        return DEFAULT_PROFILES

    def save_profiles(self):
        with open(PROFILES_FILE, 'w') as f:
            json.dump(self.profiles, f, indent=4)

    # --- License Management Methods ---
    def load_license_data(self):
        """Load license data from file"""
        if os.path.exists(LICENSE_FILE):
            try:
                with open(LICENSE_FILE, 'r') as f:
                    return json.load(f)
            except:
                return {}
        return {}

    def save_license_data(self, data):
        """Save license data to file"""
        with open(LICENSE_FILE, 'w') as f:
            json.dump(data, f, indent=4)
        
    def update_license_keys_file(self, license_data):
        """Update the license keys text file with all license information"""
        keys_file_path = os.path.join(DB_FOLDER, "license_keys.txt")
        
        try:
            with open(keys_file_path, 'w') as f:
                f.write("TCG TECH INVOICE GENERATOR - LICENSE KEYS\n")
                f.write("=" * 50 + "\n\n")
                
                # Write all valid license keys
                f.write("VALID LICENSE KEYS:\n")
                f.write("-" * 20 + "\n")
                current_year = str(datetime.now().year)
                f.write(f"Current Year: {current_year}\n\n")
                for account_type, key in VALID_LICENSE_KEYS.items():
                    f.write(f"{account_type.title()} Account: {key}\n")
                
                f.write("\n")
                
                # Write current active license info
                if license_data and license_data.get("is_active", False):
                    f.write("CURRENT ACTIVE LICENSE:\n")
                    f.write("-" * 25 + "\n")
                    f.write(f"Account Type: {license_data.get('account_type', 'Unknown').title()}\n")
                    f.write(f"License Key: {license_data.get('license_key', 'N/A')}\n")
                    f.write(f"Activation Date: {license_data.get('activation_date', 'N/A')}\n")
                    
                    # Calculate expiration info
                    if license_data.get('activation_date'):
                        try:
                            activation_date = datetime.fromisoformat(license_data['activation_date'])
                            current_date = datetime.now()
                            days_since_activation = (current_date - activation_date).days
                            
                            account_type = license_data.get('account_type', '')
                            if account_type == 'subscription':
                                days_remaining = max(0, 365 - days_since_activation)
                                expiration_date = activation_date.replace(year=activation_date.year + 1)
                                f.write(f"Expiration Date: {expiration_date.strftime('%Y-%m-%d')}\n")
                                f.write(f"Days Remaining: {days_remaining}\n")
                            elif account_type == 'demo':
                                days_remaining = max(0, 15 - days_since_activation)
                                expiration_date = activation_date.replace(day=activation_date.day + 15)
                                f.write(f"Expiration Date: {expiration_date.strftime('%Y-%m-%d')}\n")
                                f.write(f"Days Remaining: {days_remaining}\n")
                            elif account_type == 'permanent':
                                f.write("Expiration: Never\n")
                                f.write("Days Remaining: Unlimited\n")
                        except:
                            f.write("Expiration Info: Error calculating\n")
                    
                    f.write(f"Status: Active\n")
                else:
                    f.write("CURRENT LICENSE STATUS: No Active License\n")
                
                f.write("\n")
                f.write("LICENSE POLICY:\n")
                f.write("-" * 15 + "\n")
                f.write("- Demo Account: 2-day trial period\n")
                f.write("- Subscription Account: 365-day annual license\n")
                f.write("- Permanent Account: Lifetime license\n")
                f.write("- All licenses require valid activation keys\n")
                
        except Exception as e:
            print(f"Error updating license keys file: {e}")

    def validate_license_key(self, license_key, account_type):
        """Validate license key against expected value and ensure it contains current year"""
        expected_key = VALID_LICENSE_KEYS.get(account_type)
        
        # Check if the key matches exactly
        if license_key != expected_key:
            return False
        
        # Additionally check that the key contains the current year
        # This ensures old keys from previous years won't work
        current_year = str(datetime.now().year)
        if current_year not in license_key:
            return False
        
        return True

    def activate_license(self, license_key, account_type):
        """Activate license and save activation data"""
        if not self.validate_license_key(license_key, account_type):
            return False, "Invalid license key"

        # Create license data
        license_data = {
            "account_type": account_type,
            "license_key": license_key,
            "activation_date": datetime.now().isoformat(),
            "is_active": True
        }

        self.save_license_data(license_data)
        return True, f"{account_type.title()} account activated successfully!"

    def check_license_status(self):
        """
        Check license status and return:
        - (True, None) if valid
        - (False, "expired") if subscription expired
        - (False, "invalid") if no valid license
        """
        license_data = self.load_license_data()

        if not license_data or not license_data.get("is_active", False):
            return False, "invalid"

        account_type = license_data.get("account_type", "")
        activation_date_str = license_data.get("activation_date", "")

        if not activation_date_str:
            return False, "invalid"

        try:
            activation_date = datetime.fromisoformat(activation_date_str)
            current_date = datetime.now()
            days_since_activation = (current_date - activation_date).days

            if account_type == "permanent":
                return True, None
            elif account_type == "subscription":
                if days_since_activation >= 365:
                    return False, "expired"
                elif days_since_activation >= 363:  # 2 days before expiration
                    return True, "expiring_soon"
                else:
                    return True, None
            elif account_type == "demo":
                if days_since_activation >= 2:  # Block after 2 days
                    return False, "expired"
                elif days_since_activation >= 1:  # Warning at 1 day
                    return True, "demo_expiring"
                else:
                    return True, None
            else:
                return False, "invalid"

        except (ValueError, TypeError):
            return False, "invalid"

    def get_license_info(self):
        """Get current license information"""
        license_data = self.load_license_data()
        if not license_data:
            return None

        account_type = license_data.get("account_type", "unknown")
        activation_date_str = license_data.get("activation_date", "")

        if activation_date_str:
            try:
                activation_date = datetime.fromisoformat(activation_date_str)
                days_remaining = None

                if account_type == "subscription":
                    current_date = datetime.now()
                    days_since_activation = (current_date - activation_date).days
                    days_remaining = max(0, 365 - days_since_activation)
                elif account_type == "demo":
                    current_date = datetime.now()
                    days_since_activation = (current_date - activation_date).days
                    days_remaining = max(0, 2 - days_since_activation)

                return {
                    "account_type": account_type,
                    "activation_date": activation_date,
                    "days_remaining": days_remaining
                }
            except:
                pass

        return None

    def show_subscription_warning(self):
        """Show subscription or demo account expiring warning"""
        license_info = self.get_license_info()
        if not license_info or license_info["days_remaining"] is None:
            return

        account_type = license_info["account_type"]
        days_remaining = license_info["days_remaining"]

        if account_type == "subscription":
            if days_remaining <= 2:
                message = f"?? WARNING: Your subscription expires in {days_remaining} day{'s' if days_remaining != 1 else ''}!\n\nPlease renew your license to continue using the application."
                messagebox.showwarning("Subscription Expiring", message)
        elif account_type == "demo":
            if days_remaining <= 1:  # Warning at 14 days (15-14=1 day remaining)
                message = "?? WARNING: Your demo account license has been going to expire!\n\nYou have limited time remaining. Please upgrade to a full account."
                messagebox.showwarning("Demo Account Expiring", message)

    def setup_sidebar(self):
        # Modern gradient-like sidebar with shadow effect
        self.sidebar = ct.CTkFrame(self, width=300, corner_radius=0, 
                                   fg_color=("#ffffff", "#1a1a1a"),
                                   border_width=0)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        
        # Configure grid for better space management
        self.sidebar.grid_rowconfigure(5, weight=1)  # Make history section expandable

        # Add Application Logo at the top with modern styling
        logo_container = ct.CTkFrame(self.sidebar, fg_color="transparent")
        logo_container.pack(pady=(15, 10))
        
        logo_path = os.path.join(RESOURCE_DIR, "Image", "logo.png")
        if os.path.exists(logo_path):
            try:
                logo_img = Image.open(logo_path)
                logo_img = logo_img.resize((80, 80), Image.Resampling.LANCZOS)
                logo_photo = ct.CTkImage(light_image=logo_img, dark_image=logo_img, size=(80, 80))
                
                logo_label = ct.CTkLabel(logo_container, image=logo_photo, text="")
                logo_label.pack()
                
                self._app_logo = logo_photo
            except Exception as e:
                print(f"Could not load logo: {e}")
                ct.CTkLabel(logo_container, text="??", 
                           font=ct.CTkFont(size=40)).pack()
        else:
            ct.CTkLabel(logo_container, text="??", 
                       font=ct.CTkFont(size=40)).pack()

        # App title with gradient effect
        ct.CTkLabel(self.sidebar, text="INVOICE", 
                   font=ct.CTkFont(size=20, weight="bold"),
                   text_color=("#1e40af", "#60a5fa")).pack()
        ct.CTkLabel(self.sidebar, text="MANAGEMENT SYSTEM", 
                   font=ct.CTkFont(size=9, weight="bold"),
                   text_color=("#64748b", "#94a3b8")).pack(pady=(0, 10))

        # Decorative line
        ct.CTkFrame(self.sidebar, height=2, fg_color=("#e2e8f0", "#334155"),
                   corner_radius=2).pack(fill="x", padx=30, pady=8)

        # User/Business Switcher Title with icon
        header_frame = ct.CTkFrame(self.sidebar, fg_color="transparent")
        header_frame.pack(pady=(8, 10))
        
        ct.CTkLabel(header_frame, text="", 
                   font=ct.CTkFont(size=16)).pack(side="left", padx=(0, 6))
        ct.CTkLabel(header_frame, text="SELECT BUSINESS",
                   font=ct.CTkFont(size=12, weight="bold"), 
                   text_color=("#475569", "#cbd5e1")).pack(side="left")

        # Business Switcher Buttons with modern card design
        self.biz_btn_frame = ct.CTkFrame(self.sidebar, fg_color="transparent")
        self.biz_btn_frame.pack(fill="x", padx=15)

        self.biz_buttons = {}

        # Create modern card-style buttons for each profile
        for biz_id, data in self.profiles.items():
            display_text = f"{data['name']}"

            btn = ct.CTkButton(
                self.biz_btn_frame,
                text=display_text,
                command=lambda b=biz_id: self.switch_business(b),
                height=45,
                font=ct.CTkFont(size=11, weight="bold"),
                fg_color=("#f8fafc", "#2d3748"),
                text_color=("#1e293b", "#e2e8f0"),
                hover_color=("#e2e8f0", "#374151"),
                border_width=2,
                border_color=("#cbd5e1", "#4a5568"),
                corner_radius=10,
                anchor="w"
            )
            btn.pack(fill="x", pady=4)
            self.biz_buttons[biz_id] = btn

        # Decorative line
        ct.CTkFrame(self.sidebar, height=2, fg_color=("#e2e8f0", "#334155"),
                   corner_radius=1).pack(fill="x", pady=15, padx=25)

        # Search Section with modern header
        search_header = ct.CTkFrame(self.sidebar, fg_color="transparent")
        search_header.pack(pady=(15, 8))
        
        ct.CTkLabel(search_header, text="", 
                   font=ct.CTkFont(size=16)).pack(side="left", padx=(0, 6))
        ct.CTkLabel(search_header, text="SEARCH INVOICES",
                   font=ct.CTkFont(size=12, weight="bold"), 
                   text_color=("#475569", "#cbd5e1")).pack(side="left")

        # Search input
        self.search_entry = ct.CTkEntry(self.sidebar, 
                                       placeholder_text="= ?? Search by ID or Client Name...",
                                       height=35,
                                       corner_radius=10,
                                       border_width=2,
                                       border_color=("#cbd5e1", "#475569"),
                                       font=ct.CTkFont(size=11))
        self.search_entry.pack(fill="x", padx=15, pady=(0, 10))
        self.search_entry.bind("<KeyRelease>", self.on_search_change)

        # History Section with modern header
        history_header = ct.CTkFrame(self.sidebar, fg_color="transparent")
        history_header.pack(pady=(15, 8))
        
        ct.CTkLabel(history_header, text="??", 
                   font=ct.CTkFont(size=16)).pack(side="left", padx=(0, 6))
        ct.CTkLabel(history_header, text="RECENT INVOICES",
                   font=ct.CTkFont(size=12, weight="bold"), 
                   text_color=("#475569", "#cbd5e1")).pack(side="left")

        # Scrollable history list with fixed height
        self.history_list = ct.CTkScrollableFrame(self.sidebar, 
                                                  fg_color="transparent",
                                                  height=280)
        self.history_list.pack(padx=15, pady=(0, 15), fill="both", expand=True)
        self.refresh_history_ui()

    def setup_main_content(self):
        # Modern main frame with gradient background
        self.main_frame = ct.CTkScrollableFrame(self, corner_radius=0, 
                                                fg_color=("#f8fafc", "#0f172a"))
        self.main_frame.grid(row=0, column=1, sticky="nsew")

        # --- Top Bar: Business Settings with modern card design ---
        # Create draggable wrapper for settings card
        settings_wrapper = self.create_draggable_container(
            self.main_frame, "biz_settings", "Business Information"
        )
        
        settings_card = ct.CTkFrame(settings_wrapper, 
                                   fg_color=("#ffffff", "#1e293b"),
                                   corner_radius=20,
                                   border_width=2,
                                   border_color=("#e2e8f0", "#334155"))
        settings_card.pack(fill="x", pady=(0, 25))

        # Card header
        card_header = ct.CTkFrame(settings_card, fg_color="transparent")
        card_header.pack(fill="x", padx=25, pady=(20, 15))
        
        ct.CTkLabel(card_header, text="", 
                   font=ct.CTkFont(size=24)).pack(side="left", padx=(0, 10))
        ct.CTkLabel(card_header, text="", 
                   font=ct.CTkFont(size=20, weight="bold"),
                   text_color=("#1e293b", "#f1f5f9")).pack(side="left")

        # Content container
        content_frame = ct.CTkFrame(settings_card, fg_color="transparent")
        content_frame.pack(fill="x", padx=25, pady=(0, 20))

        # Logo Display with modern styling
        logo_frame = ct.CTkFrame(content_frame, 
                                fg_color=("#f8fafc", "#0f172a"),
                                corner_radius=15,
                                border_width=2,
                                border_color=("#cbd5e1", "#475569"))
        logo_frame.pack(side="left", padx=(0, 25), pady=10)
        
        self.logo_preview = ct.CTkLabel(logo_frame, text="No Logo", 
                                       width=100, height=100,
                                       fg_color=("#e2e8f0", "#334155"),
                                       corner_radius=12,
                                       font=ct.CTkFont(size=11))
        self.logo_preview.pack(padx=10, pady=10)

        # Business Info Inputs with modern styling
        info_frame = ct.CTkFrame(content_frame, fg_color="transparent")
        info_frame.pack(side="left", fill="both", expand=True, pady=10)

        self.biz_name_entry = ct.CTkEntry(info_frame, width=350, height=45,
                                          font=ct.CTkFont(size=18, weight="bold"),
                                          placeholder_text="= Business Name",
                                          corner_radius=12,
                                          border_width=2,
                                          border_color=("#cbd5e1", "#475569"))
        self.biz_name_entry.pack(anchor="w", pady=(0, 10))

        self.biz_addr_entry = ct.CTkEntry(info_frame, width=350, height=40,
                                          placeholder_text="= Business Address",
                                          corner_radius=10,
                                          border_width=2,
                                          border_color=("#cbd5e1", "#475569"))
        self.biz_addr_entry.pack(anchor="w", pady=(0, 10))

        self.biz_email_entry = ct.CTkEntry(info_frame, width=350, height=40,
                                           placeholder_text="= Business Email",
                                           corner_radius=10,
                                           border_width=2,
                                           border_color=("#cbd5e1", "#475569"))
        self.biz_email_entry.pack(anchor="w", pady=(0, 10))
        self.biz_email_entry.bind("<KeyRelease>", self.validate_business_email)
        
        # Business email validation warning label (inline - will be positioned separately)
        self.biz_email_warning_label = ct.CTkLabel(info_frame, text="", 
                                                    font=ct.CTkFont(size=10),
                                                    text_color="red")

        # Phone with inline warning
        phone_container = ct.CTkFrame(info_frame, fg_color="transparent")
        phone_container.pack(anchor="w", fill="x", pady=(0, 10))
        
        self.biz_phone_entry = ct.CTkEntry(phone_container, width=350, height=40,
                                           placeholder_text="= Business Phone",
                                           corner_radius=10,
                                           border_width=2,
                                           border_color=("#cbd5e1", "#475569"))
        self.biz_phone_entry.pack(side="left")
        self.biz_phone_entry.bind("<KeyRelease>", self.validate_phone)
        
        # Phone validation warning label (inline on right)
        self.phone_warning_label = ct.CTkLabel(phone_container, text="", 
                                               font=ct.CTkFont(size=10),
                                               text_color="red")
        self.phone_warning_label.pack(side="left", padx=(10, 0))

        # GST with inline warning
        gst_container = ct.CTkFrame(info_frame, fg_color="transparent")
        gst_container.pack(anchor="w", fill="x", pady=(0, 10))
        
        self.biz_gst_entry = ct.CTkEntry(gst_container, width=350, height=40,
                                         placeholder_text="= GST No (Optional)",
                                         corner_radius=10,
                                         border_width=2,
                                         border_color=("#cbd5e1", "#475569"))
        self.biz_gst_entry.pack(side="left")
        self.biz_gst_entry.bind("<KeyRelease>", self.validate_gst)
        
        # GST validation warning label (inline on right)
        self.gst_warning_label = ct.CTkLabel(gst_container, text="", 
                                             font=ct.CTkFont(size=10),
                                             text_color="red")
        self.gst_warning_label.pack(side="left", padx=(10, 0))

        self.biz_watermark_entry = ct.CTkEntry(info_frame, width=350, height=40,
                                               placeholder_text="= Watermark Text (Optional)",
                                               corner_radius=10,
                                               border_width=2,
                                               border_color=("#cbd5e1", "#475569"))
        self.biz_watermark_entry.pack(anchor="w", pady=(0, 10))

        # Right side frame for Upload Logo and Save Buttons
        right_frame = ct.CTkFrame(content_frame, fg_color="transparent")
        right_frame.pack(side="right", padx=20, pady=10)

        # Upload Logo and Save Buttons in the right frame
        ct.CTkButton(right_frame, text="Upload Logo", 
                    command=self.upload_logo, 
                    width=250, height=45,
                    fg_color=("#64748b", "#475569"),
                    hover_color=("#475569", "#334155"),
                    corner_radius=12,
                    font=ct.CTkFont(size=13, weight="bold")).pack(pady=(0, 8))
        
        self.save_biz_btn = ct.CTkButton(right_frame, text="Save Info", 
                                         command=self.save_current_profile_info,
                                         width=250, height=45,
                                         corner_radius=12,
                                         font=ct.CTkFont(size=13, weight="bold"))
        self.save_biz_btn.pack(pady=(0, 8))
        
        # Layout Customizer Button - REMOVED
        
        # --- Divider ---
        ct.CTkFrame(self.main_frame, height=3, 
                   fg_color=("#e2e8f0", "#334155"),
                   corner_radius=2).pack(fill="x", padx=40, pady=20)

        # --- Invoice Form with modern card design ---
        # Create draggable wrapper for invoice card
        invoice_wrapper = self.create_draggable_container(
            self.main_frame, "client_details", "Client Details"
        )
        
        invoice_card = ct.CTkFrame(invoice_wrapper,
                                  fg_color=("#ffffff", "#1e293b"),
                                  corner_radius=20,
                                  border_width=2,
                                  border_color=("#e2e8f0", "#334155"))
        invoice_card.pack(fill="x", pady=(0, 20))

        # Card header
        invoice_header = ct.CTkFrame(invoice_card, fg_color="transparent")
        invoice_header.pack(fill="x", padx=25, pady=(20, 15))
        
        ct.CTkLabel(invoice_header, text="", 
                   font=ct.CTkFont(size=22)).pack(side="left", padx=(0, 10))
        ct.CTkLabel(invoice_header, text="CLIENT DETAILS", 
                   font=ct.CTkFont(size=18, weight="bold"),
                   text_color=("#1e293b", "#f1f5f9")).pack(side="left")

        client_frame = ct.CTkFrame(invoice_card, fg_color="transparent")
        client_frame.pack(fill="x", padx=20, pady=(0, 20))
        
        # First row - Client Name and Address (same height, aligned)
        row1 = ct.CTkFrame(client_frame, fg_color="transparent")
        row1.pack(fill="x", pady=(0, 15))
        
        self.client_name = self.create_input_fixed_width(row1, "Client Name", "", width=300)
        self.client_address = self.create_input_fixed_width(row1, "Client Address", "", width=300)
        
        # Second row - Client Phone and Email (same height, aligned)
        row2 = ct.CTkFrame(client_frame, fg_color="transparent")
        row2.pack(fill="x", pady=(0, 15))
        
        self.client_phone, self.client_phone_warning = self.create_input_with_validation_fixed(row2, "Client Phone", "", width=300)
        self.client_email, self.client_email_warning = self.create_input_with_validation_fixed(row2, "Client Email", "", width=300)
        
        # Third row - Notes (full width text area with save/edit buttons)
        row3 = ct.CTkFrame(client_frame, fg_color="transparent")
        row3.pack(fill="x", pady=(0, 15))
        
        notes_container = ct.CTkFrame(row3, fg_color="transparent")
        notes_container.pack(side="left", padx=10, fill="x", expand=True)
        
        # Notes header with buttons
        notes_header = ct.CTkFrame(notes_container, fg_color="transparent")
        notes_header.pack(anchor="w", fill="x")
        
        ct.CTkLabel(notes_header, text="Notes (Optional)", 
                   font=ct.CTkFont(size=13, weight="bold"),
                   text_color=("#475569", "#cbd5e1")).pack(side="left")
        
        # Save Notes button
        ct.CTkButton(notes_header, text="💾 Save Notes", width=100, height=28,
                    command=self.save_notes_template,
                    fg_color="#10b981", hover_color="#059669").pack(side="left", padx=10)
        
        # Load Saved Notes button
        ct.CTkButton(notes_header, text="📝 Load Saved", width=100, height=28,
                    command=self.load_notes_template,
                    fg_color="#3b82f6", hover_color="#2563eb").pack(side="left")
        
        self.client_notes = ct.CTkTextbox(notes_container, height=60, width=620,
                                         font=ct.CTkFont(size=13))
        self.client_notes.pack(fill="x")
        
        # Load saved notes if exists
        self.load_notes_on_startup()
        
        # Fourth row - Invoice Date, Invoice #, Due Date (same height, aligned)
        row4 = ct.CTkFrame(client_frame, fg_color="transparent")
        row4.pack(fill="x")
        
        # Bind validation events for client fields
        self.client_email.bind("<KeyRelease>", self.validate_client_email)
        self.client_phone.bind("<KeyRelease>", self.validate_client_phone)
        
        # Invoice Date
        invoice_date_container = ct.CTkFrame(row4, fg_color="transparent")
        invoice_date_container.pack(side="left", padx=10)
        
        label_frame = ct.CTkFrame(invoice_date_container, fg_color="transparent")
        label_frame.pack(anchor="w", fill="x")
        
        lbl = ct.CTkLabel(label_frame, text="Invoice Date", 
                         font=ct.CTkFont(size=13, weight="bold"),
                         text_color=("#475569", "#cbd5e1"))
        lbl.pack(side="left")
        
        self.invoice_date = ct.CTkEntry(invoice_date_container, height=40, width=200,
                                        corner_radius=10,
                                        border_width=2,
                                        border_color=("#cbd5e1", "#475569"),
                                        font=ct.CTkFont(size=12))
        self.invoice_date.pack(fill="x", pady=(5, 0))
        # Set today's date as default
        self.invoice_date.insert(0, datetime.now().strftime("%d-%m-%Y"))
        
        # Invoice Number
        invoice_num_container = ct.CTkFrame(row4, fg_color="transparent")
        invoice_num_container.pack(side="left", padx=10)
        
        label_frame2 = ct.CTkFrame(invoice_num_container, fg_color="transparent")
        label_frame2.pack(anchor="w", fill="x")
        
        lbl2 = ct.CTkLabel(label_frame2, text="Invoice #", 
                         font=ct.CTkFont(size=13, weight="bold"),
                         text_color=("#475569", "#cbd5e1"))
        lbl2.pack(side="left")
        
        self.invoice_num = ct.CTkEntry(invoice_num_container, height=40, width=200,
                                       corner_radius=10,
                                       border_width=2,
                                       border_color=("#cbd5e1", "#475569"),
                                       font=ct.CTkFont(size=12))
        self.invoice_num.pack(fill="x", pady=(5, 0))
        self.invoice_num.insert(0, self.get_next_invoice_number())
        
        # Due Date with Pending Checkbox - modern styling
        due_date_container = ct.CTkFrame(row4, fg_color="transparent")
        due_date_container.pack(side="left", padx=10)
        
        label_frame = ct.CTkFrame(due_date_container, fg_color="transparent")
        label_frame.pack(anchor="w", fill="x")
        
        lbl = ct.CTkLabel(label_frame, text="Due Date", 
                         font=ct.CTkFont(size=13, weight="bold"),
                         text_color=("#475569", "#cbd5e1"))
        lbl.pack(side="left")
        
        due_date_inner = ct.CTkFrame(due_date_container, fg_color="transparent")
        due_date_inner.pack(fill="x", pady=(5, 0))
        
        self.due_date = ct.CTkEntry(due_date_inner, height=40,
                                    corner_radius=10,
                                    border_width=2,
                                    border_color=("#cbd5e1", "#475569"))
        self.due_date.pack(side="left", fill="x", expand=True, padx=(0, 10))
        
        self.pending_var = ct.BooleanVar(value=False)
        self.pending_checkbox = ct.CTkCheckBox(
            due_date_inner, 
            text="? Pending", 
            variable=self.pending_var,
            command=self.toggle_pending,
            font=ct.CTkFont(size=12, weight="bold"),
            corner_radius=8
        )
        self.pending_checkbox.pack(side="left")

        # --- Items Section with modern card ---
        # Create draggable wrapper for items card
        items_wrapper = self.create_draggable_container(
            self.main_frame, "invoice_items", "Invoice Items"
        )
        
        items_card = ct.CTkFrame(items_wrapper,
                                fg_color=("#ffffff", "#1e293b"),
                                corner_radius=20,
                                border_width=2,
                                border_color=("#e2e8f0", "#334155"))
        items_card.pack(fill="x", pady=(0, 20))

        # Card header
        items_header = ct.CTkFrame(items_card, fg_color="transparent")
        items_header.pack(fill="x", padx=25, pady=(20, 15))
        
        ct.CTkLabel(items_header, text="??", 
                   font=ct.CTkFont(size=22)).pack(side="left", padx=(0, 10))
        ct.CTkLabel(items_header, text="INVOICE ITEMS", 
                   font=ct.CTkFont(size=18, weight="bold"),
                   text_color=("#1e293b", "#f1f5f9")).pack(side="left")

        entry_grid = ct.CTkFrame(items_card, fg_color="transparent")
        entry_grid.pack(fill="x", padx=25, pady=(0, 15))

        self.item_desc = ct.CTkEntry(entry_grid, placeholder_text="Item Description", 
                                     width=400, height=45,
                                     corner_radius=12,
                                     border_width=2,
                                     font=ct.CTkFont(size=13))
        self.item_desc.pack(side="left", padx=5)

        self.item_price = ct.CTkEntry(entry_grid, placeholder_text="Amount (Rs )", 
                                      width=150, height=45,
                                      corner_radius=12,
                                      border_width=2,
                                      font=ct.CTkFont(size=13))
        self.item_price.pack(side="left", padx=5)

        # Quantity checkbox and fields
        self.quantity_var = ct.BooleanVar(value=False)
        self.quantity_checkbox = ct.CTkCheckBox(
            entry_grid,
            text="Quantity",
            variable=self.quantity_var,
            command=self.toggle_quantity_fields,
            font=ct.CTkFont(size=12, weight="bold"),
            corner_radius=8
        )
        self.quantity_checkbox.pack(side="left", padx=5)

        # Quantity input (initially hidden)
        self.item_quantity = ct.CTkEntry(entry_grid, placeholder_text="e.g. 1/2, 15.5",
                                        width=100, height=45,
                                        corner_radius=12,
                                        border_width=2,
                                        font=ct.CTkFont(size=13))
        
        # Unit dropdown (initially hidden)
        self.item_unit = ct.CTkComboBox(entry_grid,
                                       values=["gm", "Kg", "Litre", "Rounds", "Trips", "Piece", "Custom"],
                                       width=100, height=45,
                                       corner_radius=12,
                                       font=ct.CTkFont(size=12),
                                       command=self.on_unit_change)
        
        # Custom unit entry (initially hidden)
        self.item_custom_unit = ct.CTkEntry(entry_grid, placeholder_text="Custom Unit",
                                           width=120, height=45,
                                           corner_radius=12,
                                           border_width=2,
                                           font=ct.CTkFont(size=13))

        self.add_btn = ct.CTkButton(entry_grid, text="? Add", 
                                    width=100, height=45,
                                    command=self.add_item,
                                    corner_radius=12,
                                    font=ct.CTkFont(size=14, weight="bold"))
        self.add_btn.pack(side="left", padx=5)

        # --- Table & Tax ---
        self.table_frame = ct.CTkFrame(items_card, 
                                      fg_color=("#f8fafc", "#0f172a"),
                                      border_width=2, 
                                      border_color=("#e2e8f0", "#334155"),
                                      corner_radius=15)
        self.table_frame.pack(fill="x", padx=25, pady=(0, 15))

        # Tax Calculation Row with modern styling
        tax_frame = ct.CTkFrame(items_card, fg_color="transparent")
        tax_frame.pack(fill="x", padx=25, pady=(0, 20))

        ct.CTkLabel(tax_frame, text="Discount (%):", font=ct.CTkFont(size=13, weight="bold"), text_color=("#475569", "#cbd5e1")).pack(side="left", padx=(0, 10))
        self.discount_entry = ct.CTkEntry(tax_frame, width=100, height=40, placeholder_text="0", corner_radius=10, border_width=2, font=ct.CTkFont(size=13))
        self.discount_entry.pack(side="left", padx=5)
        self.discount_entry.insert(0, "0")
        self.discount_entry.bind("<KeyRelease>", lambda e: self.refresh_table())

        ct.CTkLabel(tax_frame, text="Tax Rate(GST) (%):", 
                   font=ct.CTkFont(size=13, weight="bold"),
                    text_color=("#475569", "#cbd5e1")).pack(side="left", padx=(20, 10))
        
        self.tax_entry = ct.CTkEntry(tax_frame, width=100, height=40,
                                     placeholder_text="0",
                                     corner_radius=10,
                                     border_width=2,
                                     font=ct.CTkFont(size=13))
        self.tax_entry.pack(side="left", padx=5)
        self.tax_entry.insert(0, "0")
        self.tax_entry.bind("<KeyRelease>", lambda e: self.refresh_table())

        ct.CTkButton(tax_frame, text="?? Update Total", 
                    width=140, height=40,
                    fg_color=("#64748b", "#475569"),
                    hover_color=("#475569", "#334155"),
                    corner_radius=10,
                    font=ct.CTkFont(size=12, weight="bold"),
                    command=self.refresh_table).pack(side="left", padx=10)

        self.refresh_table()

        # --- Bottom Actions with modern buttons ---
        # Create draggable wrapper for action buttons
        action_wrapper = self.create_draggable_container(
            self.main_frame, "action_buttons", "Action Buttons"
        )
        
        action_frame = ct.CTkFrame(action_wrapper, fg_color="transparent")
        action_frame.pack(fill="x", pady=(0, 30))

        self.btn_pdf = ct.CTkButton(action_frame, text="?? Generate PDF", 
                                    height=55,
                                    corner_radius=15,
                                    font=ct.CTkFont(size=15, weight="bold"),
                                    command=lambda: self.generate("pdf"))
        self.btn_pdf.pack(side="right", padx=10)

        self.btn_word = ct.CTkButton(action_frame, text="Generate Word", 
                                     height=55,
                                     corner_radius=15,
                                     font=ct.CTkFont(size=15, weight="bold"),
                                     command=lambda: self.generate("word"), 
                                     fg_color=("#3b82f6", "#2563eb"),
                                     hover_color=("#2563eb", "#1d4ed8"))
        self.btn_word.pack(side="right", padx=10)

        # Print button - generates PDF and opens it for printing without saving
        self.btn_print = ct.CTkButton(action_frame, text="??? Print Invoice", 
                                     height=55,
                                     corner_radius=15,
                                     font=ct.CTkFont(size=15, weight="bold"),
                                     command=self.print_invoice,
                                     fg_color=("#059669", "#047857"),
                                     hover_color=("#047857", "#065f46"))
        self.btn_print.pack(side="right", padx=10)

        ct.CTkButton(action_frame, text="Clear Form", 
                    height=55,
                    corner_radius=15,
                    font=ct.CTkFont(size=15, weight="bold"),
                    fg_color=("#64748b", "#475569"),
                    hover_color=("#475569", "#334155"),
                    command=self.clear_form).pack(side="right", padx=10)

    def create_input(self, master, placeholder, icon=""):
        container = ct.CTkFrame(master, fg_color="transparent")
        container.pack(side="left", fill="x", expand=True, padx=10)
        
        # Label with icon
        label_frame = ct.CTkFrame(container, fg_color="transparent")
        label_frame.pack(anchor="w", fill="x")
        
        if icon:
            ct.CTkLabel(label_frame, text=icon, 
                       font=ct.CTkFont(size=14)).pack(side="left", padx=(0, 5))
        
        lbl = ct.CTkLabel(label_frame, text=placeholder, 
                         font=ct.CTkFont(size=13, weight="bold"),
                         text_color=("#475569", "#cbd5e1"))
        lbl.pack(side="left")
        
        entry = ct.CTkEntry(container, height=40,
                           corner_radius=10,
                           border_width=2,
                           border_color=("#cbd5e1", "#475569"),
                           font=ct.CTkFont(size=12))
        entry.pack(fill="x", pady=(5, 0))
        return entry
    
    def create_input_fixed_width(self, master, placeholder, icon="", width=300):
        """Create input field with fixed width for better alignment"""
        container = ct.CTkFrame(master, fg_color="transparent")
        container.pack(side="left", padx=10)
        
        # Label with icon
        label_frame = ct.CTkFrame(container, fg_color="transparent")
        label_frame.pack(anchor="w", fill="x")
        
        if icon:
            ct.CTkLabel(label_frame, text=icon, 
                       font=ct.CTkFont(size=14)).pack(side="left", padx=(0, 5))
        
        lbl = ct.CTkLabel(label_frame, text=placeholder, 
                         font=ct.CTkFont(size=13, weight="bold"),
                         text_color=("#475569", "#cbd5e1"))
        lbl.pack(side="left")
        
        entry = ct.CTkEntry(container, height=40, width=width,
                           corner_radius=10,
                           border_width=2,
                           border_color=("#cbd5e1", "#475569"),
                           font=ct.CTkFont(size=12))
        entry.pack(pady=(5, 0))
        
        # Add invisible spacer to match validation label height
        spacer = ct.CTkLabel(container, text="", height=15)
        spacer.pack(anchor="w", pady=(2, 0))
        
        return entry
    
    def create_input_with_validation(self, master, placeholder, icon=""):
        """Create input field with validation warning label"""
        container = ct.CTkFrame(master, fg_color="transparent")
        container.pack(side="left", fill="x", expand=True, padx=10)
        
        # Label with icon
        label_frame = ct.CTkFrame(container, fg_color="transparent")
        label_frame.pack(anchor="w", fill="x")
        
        if icon:
            ct.CTkLabel(label_frame, text=icon, 
                       font=ct.CTkFont(size=14)).pack(side="left", padx=(0, 5))
        
        lbl = ct.CTkLabel(label_frame, text=placeholder, 
                         font=ct.CTkFont(size=13, weight="bold"),
                         text_color=("#475569", "#cbd5e1"))
        lbl.pack(side="left")
        
        entry = ct.CTkEntry(container, height=40,
                           corner_radius=10,
                           border_width=2,
                           border_color=("#cbd5e1", "#475569"),
                           font=ct.CTkFont(size=12))
        entry.pack(fill="x", pady=(5, 0))
        
        # Warning label
        warning_label = ct.CTkLabel(container, text="", 
                                    font=ct.CTkFont(size=10),
                                    text_color="red")
        warning_label.pack(anchor="w", pady=(2, 0))
        
        return entry, warning_label
    
    def create_input_with_validation_fixed(self, master, placeholder, icon="", width=300):
        """Create input field with validation warning label and fixed width"""
        container = ct.CTkFrame(master, fg_color="transparent")
        container.pack(side="left", padx=10)
        
        # Label with icon
        label_frame = ct.CTkFrame(container, fg_color="transparent")
        label_frame.pack(anchor="w", fill="x")
        
        if icon:
            ct.CTkLabel(label_frame, text=icon, 
                       font=ct.CTkFont(size=14)).pack(side="left", padx=(0, 5))
        
        lbl = ct.CTkLabel(label_frame, text=placeholder, 
                         font=ct.CTkFont(size=13, weight="bold"),
                         text_color=("#475569", "#cbd5e1"))
        lbl.pack(side="left")
        
        entry = ct.CTkEntry(container, height=40, width=width,
                           corner_radius=10,
                           border_width=2,
                           border_color=("#cbd5e1", "#475569"),
                           font=ct.CTkFont(size=12))
        entry.pack(pady=(5, 0))
        
        # Warning label with fixed height to maintain alignment
        warning_label = ct.CTkLabel(container, text="", 
                                    font=ct.CTkFont(size=10),
                                    text_color="red", width=width, height=15)
        warning_label.pack(anchor="w", pady=(2, 0))
        
        return entry, warning_label

    def toggle_pending(self):
        """Toggle the pending status and update due date field"""
        if self.pending_var.get():
            self.due_date.configure(state="disabled", placeholder_text="PENDING")
            self.due_date.delete(0, 'end')
        else:
            self.due_date.configure(state="normal", placeholder_text="")

    def toggle_quantity_fields(self):
        """Toggle visibility of quantity input fields"""
        if self.quantity_var.get():
            self.item_quantity.pack(side="left", padx=5)
            self.item_unit.pack(side="left", padx=5)
            self.item_unit.set("Kg")  # Default value
        else:
            self.item_quantity.pack_forget()
            self.item_unit.pack_forget()
            self.item_custom_unit.pack_forget()

    def on_unit_change(self, choice):
        """Show custom unit entry if Custom is selected"""
        if choice == "Custom":
            self.item_custom_unit.pack(side="left", padx=5, before=self.add_btn)
        else:
            self.item_custom_unit.pack_forget()

    # --- Business Profile Logic ---

    def switch_business(self, biz_id):
        # AUTO-SAVE LOGIC: Save current input to the OLD profile before switching
        current_name = self.biz_name_entry.get()
        current_addr = self.biz_addr_entry.get()
        current_email = self.biz_email_entry.get()
        current_phone = self.biz_phone_entry.get()
        current_gst = self.biz_gst_entry.get().strip()
        current_watermark = self.biz_watermark_entry.get().strip()

        if current_name:
            self.profiles[self.current_biz_id]["name"] = current_name
            self.profiles[self.current_biz_id]["address"] = current_addr
            self.profiles[self.current_biz_id]["email"] = current_email
            self.profiles[self.current_biz_id]["phone"] = current_phone
            self.profiles[self.current_biz_id]["gst_no"] = current_gst
            self.profiles[self.current_biz_id]["watermark"] = current_watermark
            self.save_profiles()  # Commit to file so it persists

        # Now switch
        self.current_biz_id = biz_id
        self.current_theme = self.profiles[biz_id]["color"]
        self.refresh_theme()
        
        # Refresh history to show current business invoices
        self.refresh_history_ui()

    def refresh_theme(self):
        # 1. Update Sidebar Buttons with modern styling
        for bid, btn in self.biz_buttons.items():
            if bid == self.current_biz_id:
                btn.configure(
                    fg_color=self.current_theme, 
                    text_color="white", 
                    border_width=0,
                    hover_color=self.current_theme
                )
            else:
                btn.configure(
                    fg_color=("#f8fafc", "#2d3748"), 
                    text_color=("#1e293b", "#e2e8f0"), 
                    border_width=2,
                    border_color=("#cbd5e1", "#4a5568"),
                    hover_color=("#e2e8f0", "#374151")
                )

        # 2. Update Inputs with current profile data
        profile = self.profiles[self.current_biz_id]

        self.biz_name_entry.delete(0, 'end')
        self.biz_name_entry.insert(0, profile["name"])

        self.biz_addr_entry.delete(0, 'end')
        self.biz_addr_entry.insert(0, profile["address"])

        self.biz_email_entry.delete(0, 'end')
        self.biz_email_entry.insert(0, profile.get("email", ""))

        self.biz_phone_entry.delete(0, 'end')
        self.biz_phone_entry.insert(0, profile.get("phone", ""))

        self.biz_gst_entry.delete(0, 'end')
        self.biz_gst_entry.insert(0, profile.get("gst_no", ""))

        self.biz_watermark_entry.delete(0, 'end')
        self.biz_watermark_entry.insert(0, profile.get("watermark", ""))

        # 3. Update Buttons Color
        self.save_biz_btn.configure(fg_color=self.current_theme)
        self.add_btn.configure(fg_color=self.current_theme)
        self.btn_pdf.configure(fg_color=self.current_theme)
        self.btn_word.configure(fg_color=self.current_theme)

        # 4. Load Logo
        self.display_logo(profile["logo"])

    def display_logo(self, path):
        if path and os.path.exists(path):
            try:
                img = Image.open(path)
                # Resize keeping aspect ratio
                base_width = 80
                w_percent = (base_width / float(img.size[0]))
                h_size = int((float(img.size[1]) * float(w_percent)))
                img = img.resize((base_width, h_size), Image.Resampling.LANCZOS)

                photo_img = ct.CTkImage(light_image=img, dark_image=img, size=(base_width, h_size))
                # Store reference to prevent garbage collection
                self.logo_preview.current_image = photo_img
                self.logo_preview.configure(image=photo_img, text="")
            except Exception as e:
                # Clear any existing image and set error text
                if hasattr(self.logo_preview, 'current_image'):
                    self.logo_preview.current_image = None
                self.logo_preview.configure(image=None, text="Error")
        else:
            # Clear any existing image and set no logo text
            if hasattr(self.logo_preview, 'current_image'):
                self.logo_preview.current_image = None
            self.logo_preview.configure(image=None, text="No Logo")

    def upload_logo(self):
        file_path = filedialog.askopenfilename(filetypes=[("Images", "*.png;*.jpg;*.jpeg")])
        if file_path:
            # Generate a safe filename inside DB/logos
            ext = os.path.splitext(file_path)[1]
            new_filename = f"{self.current_biz_id}_logo{ext}"
            dest_path = os.path.join(LOGOS_FOLDER, new_filename)

            try:
                shutil.copy(file_path, dest_path)
                # Update Profile
                self.profiles[self.current_biz_id]["logo"] = dest_path
                self.save_profiles()
                self.display_logo(dest_path)
                
                # Force reload the profile to ensure logo is used
                self.refresh_theme()
                
                messagebox.showinfo("Success", "Logo updated for this business profile.")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save logo: {e}")

    def validate_business_email(self, event=None):
        """Validate business email format"""
        import re
        email = self.biz_email_entry.get().strip()
        
        # Simple email regex pattern
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        
        if email and not re.match(email_pattern, email):
            self.biz_email_warning_label.configure(text="Invalid email format")
            self.biz_email_entry.configure(border_color="red")
        else:
            self.biz_email_warning_label.configure(text="")
            self.biz_email_entry.configure(border_color=("#cbd5e1", "#475569"))
    
    def validate_phone(self, event=None):
        """Validate phone number - should be 10 digits"""
        phone = self.biz_phone_entry.get().strip()
        
        # Remove any non-digit characters for validation
        digits_only = ''.join(filter(str.isdigit, phone))
        
        if phone and len(digits_only) != 10:
            self.phone_warning_label.configure(text="Phone number must be exactly 10 digits")
            self.biz_phone_entry.configure(border_color="red")
        else:
            self.phone_warning_label.configure(text="")
            self.biz_phone_entry.configure(border_color=("#cbd5e1", "#475569"))
    
    def validate_gst(self, event=None):
        """Validate GST number - should be 15 characters"""
        gst = self.biz_gst_entry.get().strip()
        
        if gst and len(gst) != 15:
            self.gst_warning_label.configure(text="GST number must be exactly 15 characters")
            self.biz_gst_entry.configure(border_color="red")
        else:
            self.gst_warning_label.configure(text="")
            self.biz_gst_entry.configure(border_color=("#cbd5e1", "#475569"))
    
    def validate_client_email(self, event=None):
        """Validate client email format"""
        import re
        email = self.client_email.get().strip()
        
        # Simple email regex pattern
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        
        if email and not re.match(email_pattern, email):
            self.client_email_warning.configure(text="Invalid email format")
            self.client_email.configure(border_color="red")
        else:
            self.client_email_warning.configure(text="")
            self.client_email.configure(border_color=("#cbd5e1", "#475569"))
    
    def validate_client_phone(self, event=None):
        """Validate client phone number - should be 10 digits"""
        phone = self.client_phone.get().strip()
        
        # Remove any non-digit characters for validation
        digits_only = ''.join(filter(str.isdigit, phone))
        
        if phone and len(digits_only) != 10:
            self.client_phone_warning.configure(text="Phone must be 10 digits")
            self.client_phone.configure(border_color="red")
        else:
            self.client_phone_warning.configure(text="")
            self.client_phone.configure(border_color=("#cbd5e1", "#475569"))
    
    def save_current_profile_info(self):
        import re
        name = self.biz_name_entry.get()
        addr = self.biz_addr_entry.get()
        email = self.biz_email_entry.get()
        phone = self.biz_phone_entry.get()
        gst_no = self.biz_gst_entry.get().strip()
        watermark = self.biz_watermark_entry.get().strip()
        
        # Validate business email
        if email:
            email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
            if not re.match(email_pattern, email):
                messagebox.showerror("Validation Error", "Business email format is invalid!")
                return
        
        # Validate business phone number before saving
        if phone:
            digits_only = ''.join(filter(str.isdigit, phone))
            if len(digits_only) != 10:
                messagebox.showerror("Validation Error", "Business phone number must be exactly 10 digits!")
                return
        
        # Validate GST number
        if gst_no and len(gst_no) != 15:
            messagebox.showerror("Validation Error", "GST number must be exactly 15 characters!")
            return
        
        if name:
            self.profiles[self.current_biz_id]["name"] = name
            self.profiles[self.current_biz_id]["address"] = addr
            self.profiles[self.current_biz_id]["email"] = email
            self.profiles[self.current_biz_id]["phone"] = phone
            self.profiles[self.current_biz_id]["gst_no"] = gst_no
            self.profiles[self.current_biz_id]["watermark"] = watermark
            self.save_profiles()
            
            # Update the sidebar button text immediately
            self.biz_buttons[self.current_biz_id].configure(
                text=f"{name}"
            )
            
            messagebox.showinfo("Saved", f"Details saved for {name}")

    # --- Core Invoice Logic ---

    def get_next_invoice_number(self):
        # Get the last invoice number for current business
        # Don't increment here - only show the next number
        last_num = self.profiles[self.current_biz_id].get("last_invoice_num", 1000)
        next_num = last_num + 1
        return str(next_num)

    def parse_quantity(self, qty_str):
        """Parse quantity string that can be fraction (1/2, 1/4) or decimal (15.5)"""
        qty_str = qty_str.strip()
        if not qty_str:
            return None, None
        
        # Check if it's a fraction (e.g., 1/2, 1/4, 3/4)
        if '/' in qty_str:
            try:
                parts = qty_str.split('/')
                if len(parts) == 2:
                    numerator = float(parts[0].strip())
                    denominator = float(parts[1].strip())
                    if denominator != 0:
                        decimal_value = numerator / denominator
                        # Return both the original fraction string and decimal value
                        return qty_str, decimal_value
            except:
                pass
        
        # Try to parse as float/decimal
        try:
            decimal_value = float(qty_str)
            return qty_str, decimal_value
        except:
            return None, None

    def add_item(self):
        try:
            desc = self.item_desc.get().strip()
            price_per_unit = float(self.item_price.get())
            
            # Get quantity data if checkbox is checked
            quantity = None
            quantity_display = None
            unit = None
            total_price = price_per_unit  # Default to single unit price
            
            if self.quantity_var.get():
                qty_value = self.item_quantity.get().strip()
                if qty_value:
                    quantity_display, quantity = self.parse_quantity(qty_value)
                    if quantity is None:
                        messagebox.showerror("Error", "Invalid quantity format. Use numbers (15.5), fractions (1/2), or decimals.")
                        return
                    unit = self.item_unit.get()
                    if unit == "Custom":
                        unit = self.item_custom_unit.get().strip() or "units"
                    
                    # Auto-calculate total price: quantity * price_per_unit
                    total_price = quantity * price_per_unit
            
            if desc:
                # Check if item with same description already exists
                item_exists = False
                for item in self.items:
                    if item['desc'].lower() == desc.lower():
                        # Update existing item
                        item['price'] = total_price  # Store calculated total price
                        item['price_per_unit'] = price_per_unit  # Store unit price for reference
                        item['quantity'] = quantity
                        item['quantity_display'] = quantity_display
                        item['unit'] = unit
                        item_exists = True
                        break
                
                # If item doesn't exist, add new one
                if not item_exists:
                    self.items.append({
                        "desc": desc, 
                        "price": total_price,  # Store calculated total price
                        "price_per_unit": price_per_unit,  # Store unit price for reference
                        "quantity": quantity,
                        "quantity_display": quantity_display,
                        "unit": unit
                    })
                
                # Clear input fields
                self.item_desc.delete(0, 'end')
                self.item_price.delete(0, 'end')
                self.item_quantity.delete(0, 'end')
                self.item_custom_unit.delete(0, 'end')
                self.refresh_table()
        except ValueError:
            messagebox.showerror("Error", "Price must be a valid number")

    def delete_item(self, item_desc):
        """Delete an item from the invoice"""
        self.items = [item for item in self.items if item['desc'] != item_desc]
        self.refresh_table()

    def edit_item(self, item_desc, item_price_per_unit):
        """Load item into text boxes for editing"""
        self.item_desc.delete(0, 'end')
        self.item_desc.insert(0, item_desc)
        self.item_price.delete(0, 'end')
        self.item_price.insert(0, str(item_price_per_unit))

    def refresh_table(self):
        for w in self.table_frame.winfo_children(): w.destroy()

        # Check if any item has quantity
        has_quantity = any(item.get('quantity') is not None for item in self.items)

        # Modern Header with gradient
        h = ct.CTkFrame(self.table_frame, 
                       fg_color=("#3b82f6", "#2563eb"),
                       corner_radius=10)
        h.pack(fill="x", padx=10, pady=10)
        
        header_content = ct.CTkFrame(h, fg_color="transparent")
        header_content.pack(fill="x", padx=15, pady=12)
        
        ct.CTkLabel(header_content, text="?? Item Description", 
                   font=ct.CTkFont(size=13, weight="bold"),
                   text_color="white").pack(side="left", padx=10)
        
        if has_quantity:
            ct.CTkLabel(header_content, text="Quantity", 
                       font=ct.CTkFont(size=13, weight="bold"),
                       text_color="white").pack(side="left", padx=(200, 10))
        
        ct.CTkLabel(header_content, text="Price (Rs)", 
                   font=ct.CTkFont(size=13, weight="bold"),
                   text_color="white").pack(side="right", padx=(0, 80))

        subtotal = 0.0
        for idx, item in enumerate(self.items):
            # Alternating row colors
            row_color = ("#f8fafc", "#1a1f2e") if idx % 2 == 0 else ("#ffffff", "#0f172a")
            
            row = ct.CTkFrame(self.table_frame, fg_color=row_color,
                            corner_radius=8)
            row.pack(fill="x", pady=3, padx=10)
            
            row_content = ct.CTkFrame(row, fg_color="transparent")
            row_content.pack(fill="x", padx=15, pady=10)
            
            # Description on left
            ct.CTkLabel(row_content, text=item['desc'],
                       font=ct.CTkFont(size=12),
                       text_color=("#1e293b", "#e2e8f0")).pack(side="left", padx=10)
            
            # Quantity in middle (if exists)
            if has_quantity:
                qty_text = ""
                if item.get('quantity') is not None:
                    # Use quantity_display if available (shows fractions), otherwise use quantity
                    display_qty = item.get('quantity_display', item['quantity'])
                    qty_text = f"{display_qty} {item.get('unit', '')}"
                ct.CTkLabel(row_content, text=qty_text,
                           font=ct.CTkFont(size=11),
                           text_color=("#64748b", "#94a3b8"),
                           width=150).pack(side="left", padx=(150, 10))
            
            # Buttons on right
            button_frame = ct.CTkFrame(row_content, fg_color="transparent")
            button_frame.pack(side="right", padx=10)
            
            # Price label
            ct.CTkLabel(button_frame, text=f"Rs {item['price']:.2f}",
                       font=ct.CTkFont(size=12, weight="bold"),
                       text_color=("#059669", "#10b981")).pack(side="left", padx=(0, 10))
            
            # Edit button
            edit_btn = ct.CTkButton(
                button_frame,
                text="??",
                command=lambda d=item['desc'], p=item.get('price_per_unit', item['price']): self.edit_item(d, p),
                width=35,
                height=30,
                fg_color=("#3b82f6", "#2563eb"),
                hover_color=("#2563eb", "#1d4ed8"),
                corner_radius=8,
                font=ct.CTkFont(size=12)
            )
            edit_btn.pack(side="left", padx=2)
            
            # Delete button
            delete_btn = ct.CTkButton(
                button_frame,
                text="X",
                command=lambda d=item['desc']: self.delete_item(d),
                width=35,
                height=30,
                fg_color=("#ef4444", "#dc2626"),
                hover_color=("#dc2626", "#b91c1c"),
                corner_radius=8,
                font=ct.CTkFont(size=12)
            )
            delete_btn.pack(side="left", padx=2)
            
            subtotal += item['price']

        # Tax Calc
        try:
            tax_rate = float(self.tax_entry.get())
        except:
            tax_rate = 0.0

        try:
            discount_rate = float(self.discount_entry.get())
        except:
            discount_rate = 0.0

        discount_amt = subtotal * (discount_rate / 100)
        taxable_subtotal = subtotal - discount_amt
        tax_amt = taxable_subtotal * (tax_rate / 100)
        grand_total = taxable_subtotal + tax_amt

        # Modern Summary Frame
        summary_frame = ct.CTkFrame(self.table_frame, 
                                   fg_color=("#f1f5f9", "#1e293b"),
                                   corner_radius=12)
        summary_frame.pack(fill="x", pady=15, padx=10)

        summary_content = ct.CTkFrame(summary_frame, fg_color="transparent")
        summary_content.pack(fill="x", padx=20, pady=15)

        # Subtotal
        subtotal_row = ct.CTkFrame(summary_content, fg_color="transparent")
        subtotal_row.pack(fill="x", pady=3)
        ct.CTkLabel(subtotal_row, text="Subtotal:", 
                   font=ct.CTkFont(size=12),
                   text_color=("#64748b", "#94a3b8")).pack(side="left")
        ct.CTkLabel(subtotal_row, text=f"Rs {subtotal:.2f}",
                   font=ct.CTkFont(size=12, weight="bold"),
                   text_color=("#1e293b", "#f1f5f9")).pack(side="right")

        # Discount
        disc_row = ct.CTkFrame(summary_content, fg_color="transparent")
        disc_row.pack(fill="x", pady=3)
        ct.CTkLabel(disc_row, text=f"Discount ({discount_rate}%):", 
                   font=ct.CTkFont(size=12),
                   text_color=("#64748b", "#94a3b8")).pack(side="left")
        ct.CTkLabel(disc_row, text=f"Rs {discount_amt:.2f}",
                   font=ct.CTkFont(size=12, weight="bold"),
                   text_color=("#1e293b", "#f1f5f9")).pack(side="right")

        # Tax
        tax_row = ct.CTkFrame(summary_content, fg_color="transparent")
        tax_row.pack(fill="x", pady=3)
        ct.CTkLabel(tax_row, text=f"Tax(GST) ({tax_rate}%):", 
                   font=ct.CTkFont(size=12),
                   text_color=("#64748b", "#94a3b8")).pack(side="left")
        ct.CTkLabel(tax_row, text=f"Rs {tax_amt:.2f}",
                   font=ct.CTkFont(size=12, weight="bold"),
                   text_color=("#1e293b", "#f1f5f9")).pack(side="right")

        # Divider
        ct.CTkFrame(summary_content, height=2, 
                   fg_color=("#cbd5e1", "#475569")).pack(fill="x", pady=8)

        # Grand Total
        total_row = ct.CTkFrame(summary_content, fg_color="transparent")
        total_row.pack(fill="x", pady=5)
        ct.CTkLabel(total_row, text="GRAND TOTAL:", 
                   font=ct.CTkFont(size=16, weight="bold"),
                   text_color=self.current_theme).pack(side="left")
        ct.CTkLabel(total_row, text=f"Rs {grand_total:.2f}",
                   font=ct.CTkFont(size=18, weight="bold"),
                   text_color=self.current_theme).pack(side="right")

        # Save for generation
        self.current_calc = {
            "subtotal": subtotal,
            "discount_rate": discount_rate,
            "discount_amt": discount_amt,
            "tax_rate": tax_rate,
            "tax_amt": tax_amt,
            "total": grand_total
        }

    def save_notes_template(self):
        """Save current notes as template for future invoices"""
        try:
            notes_content = self.client_notes.get("1.0", "end-1c").strip()
            if not notes_content:
                messagebox.showwarning("Empty Notes", "Please enter some notes before saving.")
                return
            
            # Save to profile
            profile = self.profiles[self.current_biz_id]
            profile['saved_notes'] = notes_content
            self.save_profiles()
            
            messagebox.showinfo("Notes Saved", "Notes template saved successfully!\n\nThis will be automatically loaded for new invoices.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save notes: {e}")
    
    def load_notes_template(self):
        """Load saved notes template"""
        try:
            profile = self.profiles[self.current_biz_id]
            saved_notes = profile.get('saved_notes', '')
            
            if not saved_notes:
                messagebox.showinfo("No Saved Notes", "No saved notes template found for this business.")
                return
            
            self.client_notes.delete("1.0", "end")
            self.client_notes.insert("1.0", saved_notes)
            messagebox.showinfo("Notes Loaded", "Saved notes template loaded successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load notes: {e}")
    
    def load_notes_on_startup(self):
        """Automatically load saved notes when creating new invoice"""
        try:
            profile = self.profiles[self.current_biz_id]
            saved_notes = profile.get('saved_notes', '')
            if saved_notes:
                self.client_notes.delete("1.0", "end")
                self.client_notes.insert("1.0", saved_notes)
        except:
            pass  # Silently fail if no saved notes

    def clear_form(self):
        self.items = []
        self.client_name.delete(0, 'end')
        self.client_email.delete(0, 'end')
        self.client_phone.delete(0, 'end')
        self.client_address.delete(0, 'end')
        self.client_notes.delete("1.0", "end")
        # Reload saved notes template after clearing
        self.load_notes_on_startup()
        self.invoice_num.delete(0, 'end')
        self.invoice_num.insert(0, self.get_next_invoice_number())
        self.invoice_date.delete(0, 'end')
        self.invoice_date.insert(0, datetime.now().strftime("%d-%m-%Y"))
        self.due_date.delete(0, 'end')
        self.pending_var.set(False)
        self.due_date.configure(state="normal", placeholder_text="")
        self.tax_entry.delete(0, 'end')
        self.tax_entry.insert(0, "0")
        self.item_desc.delete(0, 'end')
        self.item_price.delete(0, 'end')
        self.refresh_table()

    # --- Generation Logic ---

    def generate(self, file_type):
        import re
        try:
            # messagebox.showinfo("Debug", f"Generate clicked for {file_type}") # DEBUG POPUP
            if not self.items:
                messagebox.showwarning("Warning", "Please add at least one item to the invoice before generating.")
                return
            
            # Validate client email
            client_email = self.client_email.get().strip()
            if client_email:
                email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
                if not re.match(email_pattern, client_email):
                    messagebox.showerror("Validation Error", "Client email format is invalid!")
                    return
            
            # Validate client phone
            client_phone = self.client_phone.get().strip()
            if client_phone:
                digits_only = ''.join(filter(str.isdigit, client_phone))
                if len(digits_only) != 10:
                    messagebox.showerror("Validation Error", "Client phone number must be exactly 10 digits!")
                    return

            self.refresh_table()
        
            # Handle pending status and due date
            due_date_value = self.due_date.get().strip() or ""
            is_pending = self.pending_var.get()
            
            # Store both values separately for better display control
            if is_pending and due_date_value:
                # Both date and pending - will be displayed on separate lines
                due_date_display = due_date_value
                pending_status = True
            elif is_pending and not due_date_value:
                # Only pending, no date
                due_date_display = ""
                pending_status = True
            else:
                # Only date, no pending
                due_date_display = due_date_value
                pending_status = False
        
            # Get invoice date from input or use today's date
            invoice_date = self.invoice_date.get() or datetime.now().strftime("%d-%m-%Y")

            # Get Profile
            profile = self.profiles[self.current_biz_id]
        
            data = {
             "id": self.invoice_num.get(),
             "date": invoice_date,
             "due_date": due_date_display,
             "is_pending": pending_status,
             "client_name": self.client_name.get(),
             "client_email": self.client_email.get(),
             "client_phone": self.client_phone.get(),
             "client_address": self.client_address.get(),
             "client_notes": self.client_notes.get("1.0", "end-1c").strip(),
             "discount_rate": self.current_calc["discount_rate"],
             "discount_amt": self.current_calc["discount_amt"],
             "items": self.items,
             "subtotal": self.current_calc["subtotal"],
             "tax_rate": self.current_calc["tax_rate"],
             "tax_amt": self.current_calc["tax_amt"],
             "total": self.current_calc["total"],
             "biz_name": profile["name"],
             "biz_addr": profile["address"],
             "biz_email": profile.get("email", ""),
             "biz_phone": profile.get("phone", ""),
             "biz_gst_no": profile.get("gst_no", ""),
             "biz_logo": profile["logo"],
             "biz_color": profile["color"],
             "biz_style": profile.get("style", "Modern"),
             "biz_watermark": profile.get("watermark", "")
            }

            # Show preview dialog first
            if not self.show_preview_dialog(data, file_type):
                return  # User cancelled
            ext = ".pdf" if file_type == "pdf" else ".docx"
            filename = f"{data['biz_name'].replace(' ', '')}_Inv{data['id']}{ext}"

            save_path = filedialog.asksaveasfilename(initialfile=filename, defaultextension=ext)
            if not save_path: return

            if file_type == "pdf":
                # Use modern style (Style 1) for all businesses
                self.make_pdf_style_1(data, save_path)
            else:
                self.make_word(data, save_path)

            self.save_to_history(data)
            
            # Increment invoice number after successful save
            invoice_num = int(data['id'])
            self.profiles[self.current_biz_id]["last_invoice_num"] = invoice_num
            self.save_profiles()
            
            # Update the invoice number field to next number
            self.invoice_num.delete(0, 'end')
            self.invoice_num.insert(0, self.get_next_invoice_number())
            
            messagebox.showinfo("Success", f"Invoice saved!")
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate invoice: {str(e)}")
            print(f"Generate Error: {e}")

    def print_invoice(self):
        """Generate PDF in temp folder and open it for printing with save option"""
        import tempfile
        import re
        
        try:
            if not self.items:
                messagebox.showwarning("Warning", "Please add at least one item to the invoice before printing.")
                return
            
            # Validate client email
            client_email = self.client_email.get().strip()
            if client_email:
                email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
                if not re.match(email_pattern, client_email):
                    messagebox.showerror("Validation Error", "Client email format is invalid!")
                    return
            
            # Validate client phone
            client_phone = self.client_phone.get().strip()
            if client_phone:
                digits_only = ''.join(filter(str.isdigit, client_phone))
                if len(digits_only) != 10:
                    messagebox.showerror("Validation Error", "Client phone number must be exactly 10 digits!")
                    return

            self.refresh_table()
        
            # Handle pending status and due date
            due_date_value = self.due_date.get().strip() or ""
            is_pending = self.pending_var.get()
            
            # Store both values separately for better display control
            if is_pending and due_date_value:
                due_date_display = due_date_value
                pending_status = True
            elif is_pending and not due_date_value:
                due_date_display = ""
                pending_status = True
            else:
                due_date_display = due_date_value
                pending_status = False
        
            # Get invoice date from input or use today's date
            invoice_date = self.invoice_date.get() or datetime.now().strftime("%d-%m-%Y")

            # Get Profile
            profile = self.profiles[self.current_biz_id]
        
            data = {
             "id": self.invoice_num.get(),
             "date": invoice_date,
             "due_date": due_date_display,
             "is_pending": pending_status,
             "client_name": self.client_name.get(),
             "client_email": self.client_email.get(),
             "client_phone": self.client_phone.get(),
             "client_address": self.client_address.get(),
             "client_notes": self.client_notes.get("1.0", "end-1c").strip(),
             "discount_rate": self.current_calc["discount_rate"],
             "discount_amt": self.current_calc["discount_amt"],
             "items": self.items,
             "subtotal": self.current_calc["subtotal"],
             "tax_rate": self.current_calc["tax_rate"],
             "tax_amt": self.current_calc["tax_amt"],
             "total": self.current_calc["total"],
             "biz_name": profile["name"],
             "biz_addr": profile["address"],
             "biz_email": profile.get("email", ""),
             "biz_phone": profile.get("phone", ""),
             "biz_gst_no": profile.get("gst_no", ""),
             "biz_logo": profile["logo"],
             "biz_color": profile["color"],
             "biz_style": profile.get("style", "Modern"),
             "biz_watermark": profile.get("watermark", "")
            }

            # Ask user if they want to save to history
            save_response = messagebox.askyesno(
                "Save Invoice?",
                "Do you want to save this invoice to history?\n\n"
                "� Yes: Invoice will be saved and can be accessed later\n"
                "� No: Invoice will only be printed (not saved)",
                icon='question'
            )

            # Create temporary PDF file
            temp_dir = tempfile.gettempdir()
            temp_filename = f"Invoice_{data['id']}_PRINT.pdf"
            temp_path = os.path.join(temp_dir, temp_filename)

            # Generate PDF in temp location
            self.make_pdf_style_1(data, temp_path)

            # Save to history if user chose Yes
            if save_response:
                self.save_to_history(data)
                
                # Increment invoice number after successful save
                invoice_num = int(data['id'])
                self.profiles[self.current_biz_id]["last_invoice_num"] = invoice_num
                self.save_profiles()
                
                # Update the invoice number field to next number
                self.invoice_num.delete(0, 'end')
                self.invoice_num.insert(0, self.get_next_invoice_number())
                
                save_message = "Invoice saved to history and opened for printing."
            else:
                save_message = "?? WARNING: This invoice will NOT be saved!\n\nThe printed invoice will not appear in your history."

            # Open PDF with default viewer
            try:
                if sys.platform == 'win32':
                    os.startfile(temp_path)
                elif sys.platform == 'darwin':  # macOS
                    os.system(f'open "{temp_path}"')
                else:  # linux
                    os.system(f'xdg-open "{temp_path}"')
                
                # Show appropriate message based on save choice
                if save_response:
                    messagebox.showinfo("Print Ready", f"{save_message}\n\nYou can now print from the PDF viewer.")
                else:
                    messagebox.showwarning("Print Ready - Not Saved", f"{save_message}\n\nYou can now print from the PDF viewer.")
                    
            except Exception as e:
                messagebox.showerror("Error", f"Could not open PDF viewer: {str(e)}")
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to prepare invoice for printing: {str(e)}")
            print(f"Print Error: {e}")

    def show_preview_dialog(self, data, file_type):
        """Show a preview dialog before generating the invoice with PDF-like format"""
        preview_window = ct.CTkToplevel(self)
        preview_window.title("Invoice Preview")
        preview_window.geometry("700x650")
        preview_window.resizable(True, True)
        preview_window.transient(self)
        preview_window.grab_set()

        # Configure grid for proper layout
        preview_window.grid_rowconfigure(0, weight=0)  # Title
        preview_window.grid_rowconfigure(1, weight=1)  # Content
        preview_window.grid_rowconfigure(2, weight=0)  # Buttons
        preview_window.grid_columnconfigure(0, weight=1)

        # Title at top
        title_label = ct.CTkLabel(preview_window, text="Invoice Preview", 
                    font=ct.CTkFont(size=18, weight="bold"))
        title_label.grid(row=0, column=0, pady=10, sticky="ew")

        # PDF-like Preview Frame (scrollable content)
        preview_frame = ct.CTkScrollableFrame(preview_window, fg_color="white", 
                                            border_width=2, border_color="#cccccc")
        preview_frame.grid(row=1, column=0, padx=15, pady=(0, 10), sticky="nsew")

        # Header with Logo and Business Info
        header_frame = ct.CTkFrame(preview_frame, fg_color="white")
        header_frame.pack(fill="x", padx=15, pady=15)

        # Logo on left
        if data['biz_logo'] and os.path.exists(data['biz_logo']):
            try:
                logo_img = Image.open(data['biz_logo'])
                logo_img = logo_img.resize((80, 80), Image.Resampling.LANCZOS)
                logo_photo = ct.CTkImage(light_image=logo_img, dark_image=logo_img, size=(80, 80))
                logo_label = ct.CTkLabel(header_frame, image=logo_photo, text="")
                logo_label.image = logo_photo
                logo_label.pack(side="left", padx=(0, 15))
            except Exception as e:
                print(f"Preview logo error: {e}")

        # Business info
        biz_info_frame = ct.CTkFrame(header_frame, fg_color="white")
        biz_info_frame.pack(side="left", fill="both", expand=True)

        ct.CTkLabel(biz_info_frame, text=data['biz_name'], 
                   font=ct.CTkFont(size=16, weight="bold"),
                   anchor="w").pack(fill="x")
        ct.CTkLabel(biz_info_frame, text=data['biz_addr'], 
                   font=ct.CTkFont(size=10), anchor="w").pack(fill="x")
        if data['biz_email']:
            ct.CTkLabel(biz_info_frame, text=data['biz_email'], 
                       font=ct.CTkFont(size=10), anchor="w").pack(fill="x")
        if data['biz_phone']:
            ct.CTkLabel(biz_info_frame, text=data['biz_phone'], 
                       font=ct.CTkFont(size=10), anchor="w").pack(fill="x")
        if data.get('biz_gst_no'):
            ct.CTkLabel(biz_info_frame, text=f"GST No: {data['biz_gst_no']}", 
                       font=ct.CTkFont(size=10), anchor="w").pack(fill="x")

        ct.CTkFrame(preview_frame, height=1, fg_color="#e0e0e0").pack(fill="x", padx=15, pady=10)

        # Invoice info
        invoice_header = ct.CTkFrame(preview_frame, fg_color="white")
        invoice_header.pack(fill="x", padx=15)

        ct.CTkLabel(invoice_header, text="INVOICE", 
                   font=ct.CTkFont(size=20, weight="bold"),
                   text_color=data['biz_color']).pack(anchor="w")
        ct.CTkLabel(invoice_header, text=f"Invoice #: {data['id']}", 
                   font=ct.CTkFont(size=11, weight="bold")).pack(anchor="w")
        ct.CTkLabel(invoice_header, text=f"Date: {data['date']}", 
                   font=ct.CTkFont(size=10)).pack(anchor="w")
        
        # Display due date and pending status
        if data.get('due_date') or data.get('is_pending'):
            if data.get('due_date'):
                ct.CTkLabel(invoice_header, text=f"Due Date: {data['due_date']}", 
                           font=ct.CTkFont(size=10)).pack(anchor="w")
            if data.get('is_pending'):
                ct.CTkLabel(invoice_header, text="PENDING", 
                           font=ct.CTkFont(size=10, weight="bold"),
                           text_color="red").pack(anchor="w")

        ct.CTkFrame(preview_frame, height=1, fg_color="#e0e0e0").pack(fill="x", padx=15, pady=10)

        # Bill To
        ct.CTkLabel(preview_frame, text="BILL TO:", 
                   font=ct.CTkFont(size=11, weight="bold")).pack(anchor="w", padx=15)
        
        client_frame = ct.CTkFrame(preview_frame, fg_color="#f8f8f8", corner_radius=5)
        client_frame.pack(fill="x", padx=15, pady=5)
        
        ct.CTkLabel(client_frame, text=data['client_name'], 
                   font=ct.CTkFont(size=11, weight="bold")).pack(anchor="w", padx=8, pady=2)
        if data.get('client_address'):
            ct.CTkLabel(client_frame, text=data['client_address'], 
                       font=ct.CTkFont(size=10)).pack(anchor="w", padx=8, pady=2)
        if data.get('client_phone'):
            ct.CTkLabel(client_frame, text=f"Phone: {data['client_phone']}", 
                       font=ct.CTkFont(size=10)).pack(anchor="w", padx=8, pady=2)
        if data.get('client_email'):
            ct.CTkLabel(client_frame, text=data['client_email'], 
                       font=ct.CTkFont(size=10)).pack(anchor="w", padx=8, pady=2)

        ct.CTkFrame(preview_frame, height=1, fg_color="#e0e0e0").pack(fill="x", padx=15, pady=10)

        # Items
        ct.CTkLabel(preview_frame, text="ITEMS:", 
                   font=ct.CTkFont(size=11, weight="bold")).pack(anchor="w", padx=15)
        
        # Check if any item has quantity
        has_quantity = any(item.get('quantity') is not None for item in data['items'])
        
        # Table Header
        table_header = ct.CTkFrame(preview_frame, fg_color=data['biz_color'])
        table_header.pack(fill="x", padx=15, pady=5)
        
        header_row = ct.CTkFrame(table_header, fg_color=data['biz_color'])
        header_row.pack(fill="x", padx=5, pady=5)
        
        ct.CTkLabel(header_row, text="Description", 
                   font=ct.CTkFont(size=11, weight="bold"),
                   text_color="white", width=250 if has_quantity else 350, anchor="w").pack(side="left", padx=5)
        
        if has_quantity:
            ct.CTkLabel(header_row, text="Quantity", 
                       font=ct.CTkFont(size=11, weight="bold"),
                       text_color="white", width=100, anchor="center").pack(side="left", padx=5)
        
        ct.CTkLabel(header_row, text="Amount (Rs )", 
                   font=ct.CTkFont(size=11, weight="bold"),
                   text_color="white", width=120, anchor="e").pack(side="right", padx=5)

        # Items
        for idx, item in enumerate(data['items']):
            item_bg = "#f8f8f8" if idx % 2 == 0 else "white"
            item_row = ct.CTkFrame(preview_frame, fg_color=item_bg)
            item_row.pack(fill="x", padx=15, pady=1)
            
            ct.CTkLabel(item_row, text=item['desc'], 
                       font=ct.CTkFont(size=10), width=250 if has_quantity else 350, anchor="w").pack(side="left", padx=8, pady=4)
            
            if has_quantity:
                qty_text = ""
                if item.get('quantity') is not None:
                    # Use quantity_display if available (shows fractions), otherwise use quantity
                    display_qty = item.get('quantity_display', item['quantity'])
                    qty_text = f"{display_qty} {item.get('unit', '')}"
                ct.CTkLabel(item_row, text=qty_text, 
                           font=ct.CTkFont(size=10), width=100, anchor="center").pack(side="left", padx=8, pady=4)
            
            ct.CTkLabel(item_row, text=f"Rs {item['price']:.2f}", 
                       font=ct.CTkFont(size=10), width=120, anchor="e").pack(side="right", padx=8, pady=4)

        ct.CTkFrame(preview_frame, height=1, fg_color="#e0e0e0").pack(fill="x", padx=15, pady=8)

        # Totals
        totals_frame = ct.CTkFrame(preview_frame, fg_color="white")
        totals_frame.pack(fill="x", padx=15, pady=5)

        subtotal_row = ct.CTkFrame(totals_frame, fg_color="white")
        subtotal_row.pack(fill="x", pady=1)
        ct.CTkLabel(subtotal_row, text="Subtotal:", 
                    font=ct.CTkFont(size=10), width=120, anchor="e").pack(side="right", padx=8)

        disc_row = ct.CTkFrame(totals_frame, fg_color="white")
        disc_row.pack(fill="x", pady=1)
        ct.CTkLabel(disc_row, text=f"Discount ({data['discount_rate']}%):", 
                    font=ct.CTkFont(size=10), width=350, anchor="e").pack(side="left", padx=8)
        ct.CTkLabel(disc_row, text=f"Rs {data['discount_amt']:.2f}", 
                   font=ct.CTkFont(size=10), width=350, anchor="e").pack(side="left", padx=8)
        ct.CTkLabel(subtotal_row, text=f"Rs {data['subtotal']:.2f}", 
                   font=ct.CTkFont(size=10), width=120, anchor="e").pack(side="right", padx=8)

        tax_row = ct.CTkFrame(totals_frame, fg_color="white")
        tax_row.pack(fill="x", pady=1)
        ct.CTkLabel(tax_row, text=f"Tax(GST) ({data['tax_rate']}%):", 
                   font=ct.CTkFont(size=10), width=350, anchor="e").pack(side="left", padx=8)
        ct.CTkLabel(tax_row, text=f"Rs {data['tax_amt']:.2f}", 
                   font=ct.CTkFont(size=10), width=120, anchor="e").pack(side="right", padx=8)

        ct.CTkFrame(totals_frame, height=2, fg_color=data['biz_color']).pack(fill="x", pady=5)
        
        total_row = ct.CTkFrame(totals_frame, fg_color="#f0f0f0")
        total_row.pack(fill="x", pady=5)
        ct.CTkLabel(total_row, text="GRAND TOTAL:", 
                   font=ct.CTkFont(size=13, weight="bold"),
                   text_color=data['biz_color'], width=350, anchor="e").pack(side="left", padx=8)
        ct.CTkLabel(total_row, text=f"Rs {data['total']:.2f}", 
                   font=ct.CTkFont(size=13, weight="bold"),
                   text_color=data['biz_color'], width=120, anchor="e").pack(side="right", padx=8)

        # Buttons at bottom - FIXED position
        button_frame = ct.CTkFrame(preview_window, fg_color="transparent")
        button_frame.grid(row=2, column=0, pady=15, sticky="ew")

        result = [False]

        def on_proceed():
            result[0] = True
            preview_window.destroy()

        def on_cancel():
            preview_window.destroy()

        ct.CTkButton(button_frame, text="Generate Invoice", command=on_proceed,
                    fg_color="#10b981", hover_color="#059669",
                    width=160, height=40,
                    font=ct.CTkFont(size=13, weight="bold")).pack(side="left", padx=(150, 10))
        ct.CTkButton(button_frame, text="Cancel", command=on_cancel,
                    fg_color="#ef4444", hover_color="#dc2626",
                    width=160, height=40,
                    font=ct.CTkFont(size=13, weight="bold")).pack(side="left", padx=10)

        preview_window.wait_window()
        return result[0]

    def image_to_base64(self, image_path):
        """Convert image to base64 string for PDF embedding"""
        try:
            with open(image_path, "rb") as image_file:
                return base64.b64encode(image_file.read()).decode('utf-8')
        except Exception as e:
            return None

    def make_pdf_style_1(self, data, path):
        """ Style 1: Modern Tech - Right Aligned Blue Theme """
        
        print(f"PDF Generation: Starting make_pdf_style_1")
        
        # Debug: Print all data to see what might contain Unicode
        print(f"DEBUG - Business name: {repr(data.get('biz_name', ''))}")
        print(f"DEBUG - Client name: {repr(data.get('client_name', ''))}")
        print(f"DEBUG - Items count: {len(data.get('items', []))}")
        if data.get('items'):
            for i, item in enumerate(data['items'][:2]):  # Show first 2 items
                print(f"DEBUG - Item {i}: {repr(item.get('desc', ''))}")
        
        # Sanitize only text fields (not the entire data dict to preserve other types)
        text_fields = ['biz_name', 'biz_addr', 'biz_email', 'biz_phone', 'biz_gst_no', 
                       'client_name', 'client_email', 'client_phone', 'client_address',
                       'id', 'date', 'due_date', 'watermark']
        for field in text_fields:
            if field in data and isinstance(data[field], str):
                original = data[field]
                data[field] = self.sanitize_text(data[field])
                if original != data[field]:
                    print(f"DEBUG - Sanitized {field}: {repr(original)} -> {repr(data[field])}")
        
        # Sanitize items
        if 'items' in data:
            for i, item in enumerate(data['items']):
                if 'desc' in item and isinstance(item['desc'], str):
                    original = item['desc']
                    item['desc'] = self.sanitize_text(item['desc'])
                    if original != item['desc']:
                        print(f"DEBUG - Sanitized item {i} desc: {repr(original)} -> {repr(item['desc'])}")
                if 'unit' in item and isinstance(item['unit'], str):
                    item['unit'] = self.sanitize_text(item['unit'])
        
        print(f"PDF Generation: All text sanitized for ASCII compatibility")
        
        pdf = FPDF()
        pdf.add_page()
        
        # COMPLETELY AVOID UNICODE - Use only ASCII characters
        font_name = 'Arial'
        rupee = 'Rs.'  # ASCII-safe only
        
        print("PDF Generation: Using ASCII-only approach to avoid encoding errors")

        # Color helper
        hex_color = data['biz_color'].lstrip('#')
        r, g, b = tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))

        # Logo on left side - only if exists
        if data['biz_logo'] and os.path.exists(data['biz_logo']):
            try:
                logo_path = data['biz_logo']
                # Convert PNG to JPG if needed
                if logo_path.lower().endswith('.png'):
                    from PIL import Image
                    img = Image.open(logo_path)
                    if img.mode == 'RGBA':
                        img = img.convert('RGB')
                    jpeg_path = logo_path.replace('.png', '_temp.jpg')
                    img.save(jpeg_path, 'JPEG')
                    logo_path = jpeg_path
                pdf.image(logo_path, 10, 8, 40)
            except Exception as e:
                print(f"Logo load failed: {e}")
                # Continue without logo

        pdf.set_font(font_name, 'B', 28)
        pdf.set_text_color(r, g, b)
        pdf.cell(0, 15, "INVOICE", 0, 1, 'R')

        pdf.set_font(font_name, '', 12)
        pdf.set_text_color(100)
        if data.get('biz_gst_no'):
            pdf.cell(0, 6, f"GST No: {data['biz_gst_no']}", 0, 1, 'R')
        pdf.cell(0, 6, f"#{data['id']}", 0, 1, 'R')
        
        pdf.ln(10)

        # Info Blocks - FROM section
        pdf.set_text_color(0)
        pdf.set_font(font_name, 'B', 13)
        pdf.cell(0, 8, "FROM:", 0, 1)

        # Increase font size for better readability
        pdf.set_font(font_name, '', 12)
        from_info = f"{data['biz_name']}\n{data['biz_addr']}"
        if data['biz_email']:
            from_info += f"\n{data['biz_email']}"
        if data['biz_phone']:
            from_info += f"\n{data['biz_phone']}"
        pdf.multi_cell(0, 6, from_info)
        
        pdf.ln(5)
        
        # TO section below FROM
        pdf.set_font(font_name, 'B', 13)
        pdf.cell(0, 8, "TO:", 0, 1)
        
        pdf.set_font(font_name, '', 12)
        client_info = f"{data['client_name']}"
        if data.get('client_address'):
            client_info += f"\n{data['client_address']}"
        if data.get('client_phone'):
            client_info += f"\nPhone: {data['client_phone']}"
        if data['client_email']:
            client_info += f"\n{data['client_email']}"
        pdf.multi_cell(0, 6, client_info)
        pdf.ln(10)

        # Add Date and Due Date on right side above table
        pdf.set_font(font_name, '', 11)
        pdf.set_text_color(0)
        pdf.cell(0, 6, f"Date: {data['date']}", 0, 1, 'R')
        
        # Display due date and pending status
        if data.get('due_date') or data.get('is_pending'):
            if data.get('due_date'):
                pdf.cell(0, 6, f"Due Date: {data['due_date']}", 0, 1, 'R')
            if data.get('is_pending'):
                pdf.set_text_color(255, 0, 0)  # Red color
                pdf.set_font(font_name, 'B', 11)
                pdf.cell(0, 6, "PENDING", 0, 1, 'R')
                pdf.set_font(font_name, '', 11)
                pdf.set_text_color(0)
        
        pdf.ln(5)

        # Check if any item has quantity
        has_quantity = any(item.get('quantity') is not None for item in data['items'])

        # Table with borders
        pdf.set_fill_color(r, g, b)
        pdf.set_text_color(255)
        pdf.set_font(font_name, 'B', 12)
        if has_quantity:
            pdf.cell(90, 10, " Description", 1, 0, 'L', True)
            pdf.cell(50, 10, " Quantity", 1, 0, 'C', True)
            pdf.cell(50, 10, " Price (Rs.)", 1, 1, 'R', True)
        else:
            pdf.cell(140, 10, " Description", 1, 0, 'L', True)
            pdf.cell(50, 10, " Price (Rs.)", 1, 1, 'R', True)

        # Items with borders
        pdf.set_text_color(0)
        pdf.set_font(font_name, '', 11)
        pdf.set_fill_color(245, 245, 245)
        fill = False
        for item in data['items']:
            if has_quantity:
                pdf.cell(90, 10, f" {item['desc']}", 1, 0, 'L', fill)
                qty_text = ""
                if item.get('quantity') is not None:
                    qty_text = f"{item['quantity']} {item.get('unit', '')}"
                pdf.cell(50, 10, qty_text, 1, 0, 'C', fill)
                pdf.cell(50, 10, f"Rs. {item['price']:.2f} ", 1, 1, 'R', fill)
            else:
                pdf.cell(140, 10, f" {item['desc']}", 1, 0, 'L', fill)
                pdf.cell(50, 10, f"Rs. {item['price']:.2f} ", 1, 1, 'R', fill)
            fill = not fill

        # Total section WITHOUT borders
        pdf.ln(5)
        pdf.set_text_color(0)
        pdf.set_font(font_name, '', 11)
        pdf.cell(140, 8, "Subtotal", 0, 0, 'R')
        pdf.cell(50, 8, f"Rs. {data['subtotal']:.2f}", 0, 1, 'R')
        
        pdf.cell(140, 8, f"Discount ({data['discount_rate']}%)", 0, 0, 'R')
        pdf.cell(50, 8, f"Rs. {data['discount_amt']:.2f}", 0, 1, 'R')

        pdf.cell(140, 8, f"Tax(GST) ({data['tax_rate']}%)", 0, 0, 'R')
        pdf.cell(50, 8, f"Rs. {data['tax_amt']:.2f}", 0, 1, 'R')

        pdf.set_text_color(r, g, b)
        pdf.set_font(font_name, 'B', 14)
        pdf.cell(140, 10, "GRAND TOTAL", 0, 0, 'R')
        pdf.cell(50, 10, f"Rs. {data['total']:.2f}", 0, 1, 'R')

        # Add notes below table on left side if exists
        if data.get('client_notes') and data['client_notes'].strip():
            pdf.ln(10)
            pdf.set_font(font_name, 'B', 12)
            pdf.set_text_color(0)
            pdf.cell(0, 8, "Notes:", 0, 1, 'L')
            pdf.set_font(font_name, '', 11)
            pdf.multi_cell(0, 6, data['client_notes'])

        # Add watermark slightly lower on the page if exists
        if data.get('biz_watermark') and data['biz_watermark'].strip():
            try:
                # Add some space and place watermark (don't use absolute positioning)
                pdf.ln(15)
                pdf.set_font('Arial', 'I', 10)
                pdf.set_text_color(150, 150, 150)
                # Convert to ASCII, replacing non-ASCII characters
                watermark_text = data['biz_watermark'].encode('ascii', 'ignore').decode('ascii')
                if watermark_text:
                    pdf.cell(0, 10, watermark_text, 0, 0, 'C')
            except Exception as e:
                print(f"Watermark error: {e}")
                # Skip watermark if any error occurs

        pdf.output(path)

    def make_pdf_style_2(self, data, path):
        """ Style 2: Nature/Classic - Centered, Green Theme """
        
        # Sanitize only text fields (not the entire data dict to preserve other types)
        text_fields = ['biz_name', 'biz_addr', 'biz_email', 'biz_phone', 'biz_gst_no', 
                       'client_name', 'client_email', 'client_phone', 'client_address',
                       'id', 'date', 'due_date', 'watermark']
        for field in text_fields:
            if field in data and isinstance(data[field], str):
                data[field] = self.sanitize_text(data[field])
        
        # Sanitize items
        if 'items' in data:
            for item in data['items']:
                if 'desc' in item and isinstance(item['desc'], str):
                    item['desc'] = self.sanitize_text(item['desc'])
                if 'unit' in item and isinstance(item['unit'], str):
                    item['unit'] = self.sanitize_text(item['unit'])
        
        pdf = FPDF()
        pdf.add_page()

        # Top Bar
        pdf.set_fill_color(5, 150, 105)  # Green
        pdf.rect(0, 0, 210, 40, 'F')

        # Logo Centered on left side - only if exists
        if data['biz_logo'] and os.path.exists(data['biz_logo']):
            try:
                logo_path = data['biz_logo']
                if logo_path.lower().endswith('.png'):
                    from PIL import Image
                    img = Image.open(logo_path)
                    if img.mode == 'RGBA':
                        img = img.convert('RGB')
                    jpeg_path = logo_path.replace('.png', '_temp.jpg')
                    img.save(jpeg_path, 'JPEG')
                    logo_path = jpeg_path
                pdf.image(logo_path, 15, 5, 35)
            except Exception as e:
                print(f"Logo load failed: {e}")

        pdf.set_y(45)
        pdf.set_font("Times", 'B', 20)
        pdf.cell(0, 10, data['biz_name'], 0, 1, 'C')
        pdf.set_font("Times", 'I', 10)
        pdf.cell(0, 5, data['biz_addr'], 0, 1, 'C')

        pdf.ln(10)
        pdf.line(20, pdf.get_y(), 190, pdf.get_y())
        pdf.ln(10)

        # Client Info
        pdf.set_font("Times", '', 12)
        pdf.cell(0, 6, f"Billed To: {data['client_name']}", 0, 1, 'L')
        if data.get('client_address'):
            pdf.cell(0, 6, f"Address: {data['client_address']}", 0, 1, 'L')
        if data.get('client_phone'):
            pdf.cell(0, 6, f"Phone: {data['client_phone']}", 0, 1, 'L')
        if data.get('client_email'):
            pdf.cell(0, 6, f"Email: {data['client_email']}", 0, 1, 'L')
        
        # Invoice details on right side
        if data.get('biz_gst_no'):
            pdf.cell(0, 6, f"GST No: {data['biz_gst_no']}", 0, 1, 'R')
        pdf.cell(0, 6, f"Invoice: #{data['id']}", 0, 1, 'R')
        pdf.cell(0, 6, f"Date: {data['date']}", 0, 1, 'R')

        pdf.ln(10)

        # Check if any item has quantity
        has_quantity = any(item.get('quantity') is not None for item in data['items'])

        # Table
        pdf.set_font("Times", 'B', 12)
        if has_quantity:
            pdf.cell(90, 8, "Item", 'B', 0)
            pdf.cell(40, 8, "Quantity", 'B', 0, 'C')
            pdf.cell(60, 8, "Cost (Rs)", 'B', 1, 'R')
        else:
            pdf.cell(150, 8, "Item", 'B', 0)
            pdf.cell(40, 8, "Cost (Rs)", 'B', 1, 'R')

        pdf.set_font("Times", '', 12)
        for item in data['items']:
            if has_quantity:
                pdf.cell(90, 8, item['desc'], 'B', 0)
                qty_text = ""
                if item.get('quantity') is not None:
                    qty_text = f"{item['quantity']} {item.get('unit', '')}"
                pdf.cell(40, 8, qty_text, 'B', 0, 'C')
                pdf.cell(60, 8, f"Rs {item['price']:.2f}", 'B', 1, 'R')
            else:
                pdf.cell(150, 8, item['desc'], 'B', 0)
                pdf.cell(40, 8, f"Rs {item['price']:.2f}", 'B', 1, 'R')

        pdf.ln(5)
        pdf.cell(150, 6, f"Discount ({data['discount_rate']}%)", 0, 0, 'R')
        pdf.cell(40, 6, f"Rs {data['discount_amt']:.2f}", 0, 1, 'R')
        
        pdf.cell(150, 6, f"Tax(GST) ({data['tax_rate']}%)", 0, 0, 'R')
        pdf.cell(40, 6, f"Rs {data['tax_amt']:.2f}", 0, 1, 'R')

        pdf.ln(5)
        pdf.set_font("Times", 'B', 16)
        pdf.cell(0, 10, f"Grand Total: Rs {data['total']:.2f}", 0, 1, 'R')

        pdf.output(path)

    def make_pdf_style_3(self, data, path):
        """ Style 3: Industrial - Orange, Bold Lines, Grid """
        
        # Sanitize only text fields (not the entire data dict to preserve other types)
        text_fields = ['biz_name', 'biz_addr', 'biz_email', 'biz_phone', 'biz_gst_no', 
                       'client_name', 'client_email', 'client_phone', 'client_address',
                       'id', 'date', 'due_date', 'watermark']
        for field in text_fields:
            if field in data and isinstance(data[field], str):
                data[field] = self.sanitize_text(data[field])
        
        # Sanitize items
        if 'items' in data:
            for item in data['items']:
                if 'desc' in item and isinstance(item['desc'], str):
                    item['desc'] = self.sanitize_text(item['desc'])
                if 'unit' in item and isinstance(item['unit'], str):
                    item['unit'] = self.sanitize_text(item['unit'])
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_draw_color(0)
        pdf.set_line_width(0.5)

        # Border around page
        pdf.rect(5, 5, 200, 287)

        # Header Box
        pdf.rect(5, 5, 200, 50)

        # Logo on left side - only if exists
        if data['biz_logo'] and os.path.exists(data['biz_logo']):
            try:
                logo_path = data['biz_logo']
                if logo_path.lower().endswith('.png'):
                    from PIL import Image
                    img = Image.open(logo_path)
                    if img.mode == 'RGBA':
                        img = img.convert('RGB')
                    jpeg_path = logo_path.replace('.png', '_temp.jpg')
                    img.save(jpeg_path, 'JPEG')
                    logo_path = jpeg_path
                pdf.image(logo_path, 10, 10, 30)
            except Exception as e:
                print(f"Logo load failed: {e}")

        pdf.set_xy(40, 10)
        pdf.set_font("Courier", 'B', 22)
        pdf.cell(100, 10, data['biz_name'].upper(), 0, 1)

        # Company info below name
        pdf.set_xy(40, 20)
        pdf.set_font("Courier", '', 10)
        company_info = f"{data['biz_addr']}\n{data['biz_email']}\n{data['biz_phone']}"
        pdf.multi_cell(120, 4, company_info)

        # Invoice details on right side
        pdf.set_xy(107, 62)
        details_text = "DETAILS:\n"
        if data.get('biz_gst_no'):
            details_text += f"GST No: {data['biz_gst_no']}\n"
        details_text += f"ID: {data['id']}\nDate: {data['date']}"
        pdf.multi_cell(90, 5, details_text)

        # Bill To section
        pdf.set_xy(10, 55)
        pdf.set_font("Courier", 'B', 12)
        pdf.cell(60, 5, "BILL TO:", 0, 1)
        pdf.set_font("Courier", '', 10)
        bill_to = f"{data['client_name']}\n"
        if data.get('client_address'):
            bill_to += f"{data['client_address']}\n"
        if data.get('client_phone'):
            bill_to += f"Phone: {data['client_phone']}\n"
        if data.get('client_email'):
            bill_to += f"{data['client_email']}"
        pdf.multi_cell(90, 4, bill_to)

        pdf.set_y(95)
        # Check if any item has quantity
        has_quantity = any(item.get('quantity') is not None for item in data['items'])
        
        # Items Grid
        pdf.set_fill_color(220, 220, 220)
        if has_quantity:
            pdf.cell(100, 10, "DESCRIPTION", 1, 0, 'L', True)
            pdf.cell(40, 10, "QUANTITY", 1, 0, 'C', True)
            pdf.cell(60, 10, "AMOUNT (Rs)", 1, 1, 'C', True)
        else:
            pdf.cell(160, 10, "DESCRIPTION", 1, 0, 'L', True)
            pdf.cell(40, 10, "AMOUNT (Rs)", 1, 1, 'C', True)

        for item in data['items']:
            if has_quantity:
                pdf.cell(100, 10, item['desc'], 1, 0)
                qty_text = ""
                if item.get('quantity') is not None:
                    # Use quantity_display if available (shows fractions), otherwise use quantity
                    display_qty = item.get('quantity_display', item['quantity'])
                    qty_text = f"{display_qty} {item.get('unit', '')}"
                pdf.cell(40, 10, qty_text, 1, 0, 'C')
                pdf.cell(60, 10, f"Rs {item['price']:.2f}", 1, 1, 'R')
            else:
                pdf.cell(160, 10, item['desc'], 1, 0)
                pdf.cell(40, 10, f"Rs {item['price']:.2f}", 1, 1, 'R')

        # Tax Rows in Grid
        pdf.cell(160, 10, f"DISCOUNT ({data['discount_rate']}%)", 1, 0, 'R')
        pdf.cell(40, 10, f"Rs {data['discount_amt']:.2f}", 1, 1, 'R')

        pdf.cell(160, 10, f"TAX(GST) ({data['tax_rate']}%)", 1, 0, 'R')
        pdf.cell(40, 10, f"Rs {data['tax_amt']:.2f}", 1, 1, 'R')

        pdf.cell(160, 10, "TOTAL DUE", 1, 0, 'R', True)
        pdf.cell(40, 10, f"Rs {data['total']:.2f}", 1, 1, 'R', True)

        pdf.output(path)

    def make_pdf_style_4(self, data, path):
        """ Style 4: Minimalist - Purple, Left Aligned, Clean """
        
        # Sanitize only text fields (not the entire data dict to preserve other types)
        text_fields = ['biz_name', 'biz_addr', 'biz_email', 'biz_phone', 'biz_gst_no', 
                       'client_name', 'client_email', 'client_phone', 'client_address',
                       'id', 'date', 'due_date', 'watermark']
        for field in text_fields:
            if field in data and isinstance(data[field], str):
                data[field] = self.sanitize_text(data[field])
        
        # Sanitize items
        if 'items' in data:
            for item in data['items']:
                if 'desc' in item and isinstance(item['desc'], str):
                    item['desc'] = self.sanitize_text(item['desc'])
                if 'unit' in item and isinstance(item['unit'], str):
                    item['unit'] = self.sanitize_text(item['unit'])
        
        pdf = FPDF()
        pdf.add_page()

        # Sidebar Color Strip
        pdf.set_fill_color(124, 58, 237)  # Purple
        pdf.rect(0, 0, 10, 297, 'F')

        pdf.set_left_margin(20)
        pdf.ln(10)

        # Logo - only if exists
        if data['biz_logo'] and os.path.exists(data['biz_logo']):
            try:
                logo_path = data['biz_logo']
                if logo_path.lower().endswith('.png'):
                    from PIL import Image
                    img = Image.open(logo_path)
                    if img.mode == 'RGBA':
                        img = img.convert('RGB')
                    jpeg_path = logo_path.replace('.png', '_temp.jpg')
                    img.save(jpeg_path, 'JPEG')
                    logo_path = jpeg_path
                pdf.image(logo_path, 20, 15, 25)
                pdf.ln(25)
            except Exception as e:
                print(f"Logo load failed: {e}")
                pdf.ln(10)
        else:
            pdf.ln(10)

        pdf.set_font("Helvetica", '', 35)
        pdf.set_text_color(50)
        pdf.cell(0, 15, "invoice.", 0, 1)

        # Invoice details on right
        pdf.set_font("Helvetica", '', 9)
        pdf.set_text_color(100)
        if data.get('biz_gst_no'):
            pdf.cell(0, 4, f"GST No: {data['biz_gst_no']}", 0, 1, 'R')
        pdf.cell(0, 4, f"Invoice #{data['id']}", 0, 1, 'R')
        pdf.cell(0, 4, f"Date: {data['date']}", 0, 1, 'R')
        pdf.ln(5)

        pdf.set_font("Helvetica", '', 10)
        pdf.set_text_color(50)
        pdf.cell(0, 5, data['client_name'], 0, 1)
        if data.get('client_address'):
            pdf.cell(0, 5, data['client_address'], 0, 1)
        if data.get('client_phone'):
            pdf.cell(0, 5, f"Phone: {data['client_phone']}", 0, 1)
        if data.get('client_email'):
            pdf.cell(0, 5, data['client_email'], 0, 1)

        pdf.ln(20)

        # Check if any item has quantity
        has_quantity = any(item.get('quantity') is not None for item in data['items'])

        # Minimal List
        for item in data['items']:
            pdf.set_font("Helvetica", 'B', 12)
            if has_quantity:
                pdf.cell(90, 8, item['desc'], 0, 0)
                pdf.set_font("Helvetica", '', 10)
                qty_text = ""
                if item.get('quantity') is not None:
                    qty_text = f"{item['quantity']} {item.get('unit', '')}"
                pdf.cell(50, 8, qty_text, 0, 0, 'C')
                pdf.set_font("Helvetica", '', 12)
                pdf.cell(40, 8, f"Rs {item['price']:.2f}", 0, 1, 'R')
            else:
                pdf.cell(140, 8, item['desc'], 0, 0)
                pdf.set_font("Helvetica", '', 12)
                pdf.cell(40, 8, f"Rs {item['price']:.2f}", 0, 1, 'R')
            # Light gray line
            pdf.set_draw_color(240)
            pdf.line(20, pdf.get_y(), 190, pdf.get_y())
            pdf.ln(2)

        pdf.ln(5)
        pdf.cell(140, 8, "Subtotal", 0, 0)
        pdf.cell(40, 8, f"?{data['subtotal']:.2f}", 0, 1, 'R')
        pdf.cell(140, 8, f"Discount ({data['discount_rate']}%)", 0, 0)
        pdf.cell(40, 8, f"?{data['discount_amt']:.2f}", 0, 1, 'R')
        pdf.cell(140, 8, f"Tax(GST) ({data['tax_rate']}%)", 0, 0)
        pdf.cell(40, 8, f"?{data['tax_amt']:.2f}", 0, 1, 'R')

        pdf.ln(5)
        pdf.set_font("Helvetica", 'B', 20)
        pdf.cell(140, 10, "Total", 0, 0)
        pdf.cell(40, 10, f"?{data['total']:.2f}", 0, 1, 'R')

        pdf.output(path)

    def sanitize_text(self, text):
        """Remove or replace characters that can't be encoded in latin-1"""
        if not text:
            return ""
        # Convert to string if not already
        text = str(text)
        
        # AGGRESSIVE ASCII-only sanitization
        try:
            # Keep only ASCII characters (0-127)
            ascii_text = ''.join(char if ord(char) < 128 else '?' for char in text)
            return ascii_text
        except Exception as e:
            print(f"Sanitization error: {e}")
            # Ultimate fallback - remove all non-alphanumeric
            return ''.join(char for char in text if char.isalnum() or char in ' .-_()[]')
    
    def get_rupee_symbol(self, format_type='pdf'):
        """
        Get the rupee symbol using Unicode code point 8377
        This is more reliable than font-dependent methods
        """
        try:
            return chr(8377)  # Unicode code point for ? symbol
        except:
            return 'Rs.'  # Fallback if Unicode fails
    
    def make_word(self, data, path):
        # Basic Word generation kept simple but robust
        doc = Document()
        
        # Use ASCII-safe rupee representation
        rupee = 'Rs.'

        # Sanitize only text fields (not the entire data dict to preserve other types)
        text_fields = ['biz_name', 'biz_addr', 'biz_email', 'biz_phone', 'biz_gst_no', 
                       'client_name', 'client_email', 'client_phone', 'client_address',
                       'id', 'date', 'due_date', 'watermark']
        for field in text_fields:
            if field in data and isinstance(data[field], str):
                data[field] = self.sanitize_text(data[field])
        
        # Sanitize items
        if 'items' in data:
            for item in data['items']:
                if 'desc' in item and isinstance(item['desc'], str):
                    item['desc'] = self.sanitize_text(item['desc'])
                if 'unit' in item and isinstance(item['unit'], str):
                    item['unit'] = self.sanitize_text(item['unit'])

        # Header Table for Logo and Biz Name
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.columns[0].width = Inches(2.0)
        header_table.columns[1].width = Inches(4.0)

        # Logo Cell
        c1 = header_table.rows[0].cells[0]
        p1 = c1.paragraphs[0]
        if data['biz_logo'] and os.path.exists(data['biz_logo']):
            run = p1.add_run()
            run.add_picture(data['biz_logo'], width=Inches(1.5))

        # Biz Info Cell
        c2 = header_table.rows[0].cells[1]
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p2.add_run(data['biz_name'] + "\n")
        run.bold = True
        run.font.size = Pt(16)
        p2.add_run(data['biz_addr'])
        if data['biz_email']:
            p2.add_run(f"\n{data['biz_email']}")
        if data['biz_phone']:
            p2.add_run(f"\n{data['biz_phone']}")
        if data.get('biz_gst_no'):
            p2.add_run(f"\nGST No: {data['biz_gst_no']}")

        doc.add_heading("INVOICE", 0)
        invoice_info = f"Invoice #: {data['id']}"
        doc.add_paragraph(invoice_info)
        
        # Add FROM section
        doc.add_heading('From:', level=2)
        from_info = f"{data['biz_name']}\n{data['biz_addr']}"
        if data['biz_email']:
            from_info += f"\n{data['biz_email']}"
        if data['biz_phone']:
            from_info += f"\n{data['biz_phone']}"
        doc.add_paragraph(from_info)
        
        # Add TO section below FROM
        doc.add_heading('To:', level=2)
        client_info = f"{data['client_name']}"
        if data.get('client_address'):
            client_info += f"\n{data['client_address']}"
        if data.get('client_phone'):
            client_info += f"\nPhone: {data['client_phone']}"
        if data.get('client_email'):
            client_info += f"\n{data['client_email']}"
        doc.add_paragraph(client_info)
        
        # Add date and due date
        date_info = f"Date: {data['date']}"
        if data.get('due_date'):
            date_info += f"\nDue Date: {data['due_date']}"
        doc.add_paragraph(date_info)
        
        # Add PENDING on separate line in red if pending
        if data.get('is_pending'):
            p = doc.add_paragraph()
            run = p.add_run("PENDING")
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)

        # Check if any item has quantity
        has_quantity = any(item.get('quantity') is not None for item in data['items'])

        if has_quantity:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Description'
            hdr_cells[1].text = 'Quantity'
            hdr_cells[2].text = 'Amount'
        else:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Description'
            hdr_cells[1].text = 'Amount'

        for item in data['items']:
            row_cells = table.add_row().cells
            row_cells[0].text = item['desc']
            if has_quantity:
                qty_text = ""
                if item.get('quantity') is not None:
                    # Use quantity_display if available (shows fractions), otherwise use quantity
                    display_qty = item.get('quantity_display', item['quantity'])
                    qty_text = f"{display_qty} {item.get('unit', '')}"
                row_cells[1].text = qty_text
                row_cells[2].text = f"{rupee} {item['price']:.2f}"
            else:
                row_cells[1].text = f"{rupee} {item['price']:.2f}"

        # Tax Rows in Word
        row_disc = table.add_row().cells
        row_disc[0].text = f"Discount ({data['discount_rate']}%)"
        row_disc[1].text = f"{rupee} {data['discount_amt']:.2f}"

        row_tax = table.add_row().cells
        row_tax[0].text = f"Tax(GST) ({data['tax_rate']}%)"
        row_tax[1].text = f"{rupee} {data['tax_amt']:.2f}"

        doc.add_paragraph(f"\nTOTAL: {rupee} {data['total']:.2f}", style='Heading 2').alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add watermark at bottom if exists
        if data.get('biz_watermark'):
            watermark_para = doc.add_paragraph()
            watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = watermark_para.add_run(data['biz_watermark'])
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(150, 150, 150)
        
        doc.save(path)

    def save_to_history(self, data):
        # Enhanced history with complete invoice data
        history_item = {
            "id": data['id'],
            "client": data['client_name'],
            "client_email": data['client_email'],
            "client_phone": data.get('client_phone', ''),
            "client_address": data.get('client_address', ''),
            "total": data['total'],
            "business": data['biz_name'],
            "items": data['items'],
            "subtotal": data['subtotal'],
            "discount_rate": data.get('discount_rate', 0),
            "discount_amt": data.get('discount_amt', 0),
            "tax_rate": data['tax_rate'],
            "tax_amt": data['tax_amt'],
            "date": data['date'],
            "due_date": data['due_date'],
            "biz_id": self.current_biz_id
        }
        
        # Add to current business history
        if self.current_biz_id not in self.history:
            self.history[self.current_biz_id] = []
        
        self.history[self.current_biz_id].append(history_item)
        
        # Store unlimited invoices - no limit removed
        # Keep all invoices for unlimited storage
        
        # Save to file
        with open(HISTORY_FILE, 'w') as f:
            json.dump(self.history, f, indent=4)
        
        self.refresh_history_ui()

    def load_history(self):
        if os.path.exists(HISTORY_FILE):
            try:
                with open(HISTORY_FILE, 'r') as f:
                    data = json.load(f)
                    # Ensure it's a dictionary with business IDs as keys
                    if isinstance(data, dict):
                        return data
                    # Convert old list format to new dict format
                    elif isinstance(data, list):
                        # Migrate old format to new format
                        new_format = {
                            "biz_1": [],
                            "biz_2": [],
                            "biz_3": [],
                            "biz_4": []
                        }
                        # Try to assign old invoices to businesses based on biz_id if available
                        for item in data:
                            biz_id = item.get('biz_id', 'biz_1')
                            if biz_id in new_format:
                                new_format[biz_id].append(item)
                        return new_format
                    else:
                        return {"biz_1": [], "biz_2": [], "biz_3": [], "biz_4": []}
            except:
                return {"biz_1": [], "biz_2": [], "biz_3": [], "biz_4": []}
        return {"biz_1": [], "biz_2": [], "biz_3": [], "biz_4": []}

    def load_invoice_from_history(self, history_item):
        """Load invoice data from history into the form"""
        # Switch to the correct business
        if history_item.get('biz_id'):
            self.switch_business(history_item['biz_id'])
        
        # Clear current form
        self.clear_form()
        
        # Load invoice data
        self.invoice_num.delete(0, 'end')
        self.invoice_num.insert(0, history_item['id'])
        
        # Load invoice date
        self.invoice_date.delete(0, 'end')
        self.invoice_date.insert(0, history_item.get('date', datetime.now().strftime("%d-%m-%Y")))
        
        self.client_name.delete(0, 'end')
        self.client_name.insert(0, history_item['client'])
        
        self.client_email.delete(0, 'end')
        self.client_email.insert(0, history_item.get('client_email', ''))
        
        self.client_phone.delete(0, 'end')
        self.client_phone.insert(0, history_item.get('client_phone', ''))
        
        self.client_address.delete(0, 'end')
        self.client_address.insert(0, history_item.get('client_address', ''))
        
        self.due_date.delete(0, 'end')
        due_date_val = history_item.get('due_date', '')
        if due_date_val == "PENDING":
            self.pending_var.set(True)
            self.toggle_pending()
        else:
            self.due_date.insert(0, due_date_val)
            self.pending_var.set(False)
        
        self.tax_entry.delete(0, 'end')
        self.tax_entry.insert(0, str(history_item.get('tax_rate', 0)))
        
        self.discount_entry.delete(0, 'end')
        self.discount_entry.insert(0, str(history_item.get('discount_rate', 0)))
        
        # Load items into the items list (this will show in the table)
        self.items = history_item.get('items', [])
        
        # Also populate the description and price text boxes with the first item (if exists)
        # This allows user to see the format and add more items easily
        if self.items:
            first_item = self.items[0]
            self.item_desc.delete(0, 'end')
            self.item_desc.insert(0, first_item['desc'])
            self.item_price.delete(0, 'end')
            self.item_price.insert(0, str(first_item['price']))
        
        # Refresh table to show loaded items
        self.refresh_table()

    def delete_invoice_from_history(self, invoice_id):
        """Delete an invoice from current business history with confirmation"""
        # Get current business history
        current_history = self.history.get(self.current_biz_id, [])
        
        # Find invoice details for confirmation
        invoice_to_delete = None
        for h in current_history:
            if h['id'] == invoice_id:
                invoice_to_delete = h
                break
        
        if invoice_to_delete:
            # Show confirmation dialog
            result = messagebox.askyesno(
                "Confirm Delete",
                f"Are you sure you want to delete invoice #{invoice_to_delete['id']} for {invoice_to_delete['client']}?\n\nThis action cannot be undone.",
                icon='warning'
            )
            
            if result:
                # Remove from current business history
                self.history[self.current_biz_id] = [h for h in current_history if h['id'] != invoice_id]
                
                # Save to file
                with open(HISTORY_FILE, 'w') as f:
                    json.dump(self.history, f, indent=4)
                
                self.refresh_history_ui()
                messagebox.showinfo("Success", f"Invoice #{invoice_id} has been deleted.")

    def on_search_change(self, event=None):
        """Handle search input changes"""
        self.refresh_history_ui()

    def filter_invoices(self, invoices, search_term):
        """Filter invoices based on search term"""
        if not search_term:
            return invoices
        
        search_term = search_term.lower().strip()
        filtered = []
        
        for invoice in invoices:
            # Search by invoice ID
            if search_term in str(invoice.get('id', '')).lower():
                filtered.append(invoice)
                continue
            
            # Search by client name
            if search_term in invoice.get('client', '').lower():
                filtered.append(invoice)
                continue
        
        return filtered

    def refresh_history_ui(self):
        for w in self.history_list.winfo_children(): w.destroy()
        
        # Get current business history
        current_history = self.history.get(self.current_biz_id, [])
        
        # Apply search filter if search entry exists
        if hasattr(self, 'search_entry'):
            search_term = self.search_entry.get()
            current_history = self.filter_invoices(current_history, search_term)
        
        if not current_history:
            empty_frame = ct.CTkFrame(self.history_list, 
                                     fg_color=("#f1f5f9", "#1e293b"),
                                     corner_radius=8)
            empty_frame.pack(fill="x", pady=10, padx=5)
            
            if hasattr(self, 'search_entry') and self.search_entry.get():
                ct.CTkLabel(empty_frame, text="No invoices found", 
                           text_color=("#94a3b8", "#64748b"),
                           font=ct.CTkFont(size=11)).pack(pady=12)
            else:
                ct.CTkLabel(empty_frame, text="No recent invoices", 
                           text_color=("#94a3b8", "#64748b"),
                           font=ct.CTkFont(size=11)).pack(pady=12)
            return
        
        # Show all matching invoices (unlimited) with compact modern card design
        for idx, h in enumerate(reversed(current_history)):
            # Compact card for each invoice
            item_card = ct.CTkFrame(self.history_list, 
                                   fg_color=("#ffffff", "#1e293b"),
                                   corner_radius=10,
                                   border_width=1,
                                   border_color=("#e2e8f0", "#334155"))
            item_card.pack(fill="x", pady=4, padx=3)
            
            # Main content frame
            content_frame = ct.CTkFrame(item_card, fg_color="transparent")
            content_frame.pack(fill="x", padx=8, pady=8)
            
            # Top row - Invoice number and amount
            top_row = ct.CTkFrame(content_frame, fg_color="transparent")
            top_row.pack(fill="x")
            
            ct.CTkLabel(top_row, text=f"#{h['id']}", 
                       font=ct.CTkFont(size=11, weight="bold"),
                       text_color=("#1e40af", "#60a5fa")).pack(side="left")
            
            ct.CTkLabel(top_row, text=f"Rs {h['total']:.2f}", 
                       font=ct.CTkFont(size=11, weight="bold"),
                       text_color=("#059669", "#10b981")).pack(side="right")
            
            # Middle row - Client name
            ct.CTkLabel(content_frame, text=h['client'], 
                       font=ct.CTkFont(size=10),
                       text_color=("#1e293b", "#f1f5f9")).pack(anchor="w", pady=(2, 0))
            
            # Bottom row - Date and status
            bottom_row = ct.CTkFrame(content_frame, fg_color="transparent")
            bottom_row.pack(fill="x", pady=(2, 0))
            
            if h.get('date'):
                ct.CTkLabel(bottom_row, text=f"{h['date']}", 
                           font=ct.CTkFont(size=9),
                           text_color=("#94a3b8", "#64748b")).pack(side="left")
            
            if h.get('due_date'):
                due_text = "? Pending" if h['due_date'] == "PENDING" else f"Due: {h['due_date']}"
                ct.CTkLabel(bottom_row, text=due_text, 
                           font=ct.CTkFont(size=9),
                           text_color=("#f59e0b", "#fbbf24")).pack(side="right")
            
            # Action buttons row
            button_row = ct.CTkFrame(item_card, fg_color="transparent")
            button_row.pack(fill="x", padx=8, pady=(0, 8))
            
            # Load button
            load_btn = ct.CTkButton(
                button_row,
                text="?? Load",
                command=lambda item=h: self.load_invoice_from_history(item),
                width=80,
                height=28,
                fg_color=("#3b82f6", "#2563eb"),
                hover_color=("#2563eb", "#1d4ed8"),
                corner_radius=6,
                font=ct.CTkFont(size=10, weight="bold")
            )
            load_btn.pack(side="left", padx=(0, 4))
            
            # Delete button
            delete_btn = ct.CTkButton(
                button_row,
                text="?",
                command=lambda item_id=h['id']: self.delete_invoice_from_history(item_id),
                width=28,
                height=28,
                fg_color=("#ef4444", "#dc2626"),
                hover_color=("#dc2626", "#b91c1c"),
                corner_radius=6,
                font=ct.CTkFont(size=12, weight="bold")
            )
            delete_btn.pack(side="right")


if __name__ == "__main__":
    # Create a temporary root window for splash screen and login
    root = ct.CTk()
    root.withdraw()  # Hide the main window
    
    # Show splash screen
    splash = SplashScreen(root)
    splash.focus_force()
    
    # Wait for splash screen to close
    root.wait_window(splash)
    
    # Check license status before proceeding
    # Create a temporary InvoiceApp instance just to check license
    temp_app = InvoiceApp.__new__(InvoiceApp)  # Create instance without calling __init__
    is_valid, status = temp_app.check_license_status()
    
    if not is_valid:
        if status == "expired":
            messagebox.showwarning("License Expired", 
                               "Your license has expired!\n\nPlease activate a new license to continue using the application.")
            # Show login screen for renewal
            login = LoginScreen(root)
            login.focus_force()
            root.wait_window(login)
            
            # After login, check license again
            is_valid, status = temp_app.check_license_status()
            if not is_valid:
                messagebox.showerror("License Error", "Failed to activate license. Application will exit.")
                root.destroy()
                exit(1)
        elif status == "invalid":
            # No valid license, show login screen
            login = LoginScreen(root)
            login.focus_force()
            root.wait_window(login)
            
            # After login, check license again
            is_valid, status = temp_app.check_license_status()
            if not is_valid:
                messagebox.showerror("License Error", "Failed to activate license. Application will exit.")
                root.destroy()
                exit(1)
    
    # Show subscription warning if expiring soon
    if status == "expiring_soon":
        license_info = temp_app.get_license_info()
        if license_info and license_info["days_remaining"] is not None:
            days_remaining = license_info["days_remaining"]
            messagebox.showwarning("Subscription Expiring Soon", 
                                 f"?? Your subscription expires in {days_remaining} day{'s' if days_remaining != 1 else ''}!\n\nPlease renew your license soon to continue using the application.")
    
    # Show demo account warning if expiring soon
    if status == "demo_expiring":
        messagebox.showwarning("Demo Account Expiring", 
                             "?? WARNING: Your demo account license has been going to expire!\n\nYou have limited time remaining. Please upgrade to a full account.")
    
    # Destroy temporary root and create actual app
    root.destroy()
    
    # Start the main application
    app = InvoiceApp()
    app.mainloop()


